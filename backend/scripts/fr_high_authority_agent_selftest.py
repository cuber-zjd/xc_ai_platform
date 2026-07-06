"""FineReport 高权限 Agent 自测脚本。

默认执行轻量自检，不依赖外部 MinIO/FineReport/SQL Server。
如需创建测试 SQL 表并校验样例数据，设置：
FR_AI_RUN_HIGH_AUTHORITY_E2E=1
"""

from __future__ import annotations

import asyncio
import os
import re
from io import BytesIO

from app.core.config import settings
from app.services.agent.fr_report.ai_operation_service import fr_report_ai_operation_service
from app.services.agent.fr_report.excel_analyzer import excel_analyzer
from app.services.agent.fr_report.high_authority_agent_service import fr_report_high_authority_agent_service


def _build_excel_bytes() -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "报价维护"
    ws.merge_cells("A1:D1")
    ws["A1"] = "博兴粮油豆粕、花生粕报价数据维护表"
    ws["A2"] = "单位：元/吨"
    ws["A3"] = "时间"
    ws["B3"] = "博兴粮油豆粕报价"
    ws.merge_cells("C3:D3")
    ws["C3"] = "青岛品品好花生粕报价"
    ws["A4"] = "指标"
    ws["B4"] = "43%"
    ws["C4"] = "黄曲霉<200PPB"
    ws["D4"] = "不保黄曲霉"
    ws["A5"] = "2026/5/13"
    ws["B5"] = 2980
    ws["C5"] = 3000
    ws["D5"] = 2960
    ws["E5"] = "=C5-D5"
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_docx_bytes() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("报表调整需求", level=1)
    doc.add_paragraph("数据来源 ncp_bocing_soybean_peanut_meal_price。price1 对应黄曲霉，price2 对应不保黄曲霉。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "筛选"
    table.cell(0, 1).text = "日期、市场"
    table.cell(1, 0).text = "填报"
    table.cell(1, 1).text = "允许维护 remark"
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _minimal_workbook(extra: str = "") -> str:
    return f"""<?xml version="1.0" encoding="UTF-8"?>
<WorkBook>
  <Version releaseVersion="11.0"/>
  <Report class="com.fr.report.worksheet.WorkSheet" name="sheet1">
    <TableDataMap/>
    <Table>
      <CellElementList>
        <C c="0" r="0"><O><![CDATA[标题]]></O><PrivilegeControl/><Expand/></C>
        <C c="1" r="5"><O><![CDATA[旧值]]></O><PrivilegeControl/><Expand/></C>
      </CellElementList>
    </Table>
  </Report>
  {extra}
</WorkBook>"""


async def _run_optional_sqlserver_selftest() -> None:
    if os.getenv("FR_AI_RUN_HIGH_AUTHORITY_E2E") != "1":
        print("跳过 SQL Server 集成自测：未设置 FR_AI_RUN_HIGH_AUTHORITY_E2E=1")
        return
    if not settings.FR_AI_SQLSERVER_ENABLED or not settings.FR_AI_SQLSERVER_HOST:
        raise RuntimeError("未配置 FR_AI_SQLSERVER_*，无法运行 SQL Server 集成自测")

    import pymssql

    ddl = """
IF OBJECT_ID('dbo.fr_ai_agent_test_quote', 'U') IS NOT NULL DROP TABLE dbo.fr_ai_agent_test_quote;
CREATE TABLE dbo.fr_ai_agent_test_quote (
    id INT IDENTITY(1,1) PRIMARY KEY,
    record_date DATE NOT NULL,
    price1 DECIMAL(18,2) NULL,
    price2 DECIMAL(18,2) NULL,
    market NVARCHAR(64) NULL,
    category NVARCHAR(64) NULL,
    remark NVARCHAR(200) NULL
);
INSERT INTO dbo.fr_ai_agent_test_quote(record_date, price1, price2, market, category, remark)
VALUES ('2026-05-13', 3000, 2960, N'青岛', N'花生粕', N'自测样例');
"""
    with pymssql.connect(
        server=settings.FR_AI_SQLSERVER_HOST,
        port=settings.FR_AI_SQLSERVER_PORT,
        user=settings.FR_AI_SQLSERVER_USER,
        password=settings.FR_AI_SQLSERVER_PASSWORD,
        database=settings.FR_AI_SQLSERVER_DATABASE,
        timeout=10,
        login_timeout=10,
        charset="UTF-8",
    ) as conn:
        with conn.cursor() as cursor:
            for statement in [item.strip() for item in ddl.split(";") if item.strip()]:
                cursor.execute(statement)
        conn.commit()

    schema = await fr_report_ai_operation_service._build_database_source_context(
        prompt="读取 dbo.fr_ai_agent_test_quote，price1 对应黄曲霉，price2 对应不保黄曲霉",
        structure=type("FakeStructure", (), {"datasets": []})(),
        source_xml=None,
    )
    assert schema["available"], f"未读取到测试表结构：{schema}"
    fields = {item["name"] for item in schema["schema"]["fields"]}
    assert {"record_date", "price1", "price2", "market", "category", "remark"}.issubset(fields)
    print("SQL Server 集成自测通过：已创建并读取 dbo.fr_ai_agent_test_quote")


async def main() -> None:
    excel_result = excel_analyzer.analyze(_build_excel_bytes(), "高权限Agent自测.xlsx")
    excel_text = str(excel_result.model_dump(mode="json"))
    assert "黄曲霉" in excel_text and "不保黄曲霉" in excel_text, "Excel 多层表头解析未识别关键指标"

    docx_text = fr_report_high_authority_agent_service._read_docx_text(_build_docx_bytes())
    assert "price1" in docx_text and "price2" in docx_text, "Word 需求解析未读取关键字段"

    source_xml = _minimal_workbook()
    fr_report_ai_operation_service._validate_full_cpt_xml(source_xml)
    editing_context = fr_report_high_authority_agent_service._build_workbook_editing_context(
        source_xml.replace(
            "<Table>",
            '<ReportPageAttr><HC F="0" T="0"/><HR F="1" T="2"/></ReportPageAttr><ColumnWidth defaultValue="2743200"><![CDATA[2743200,2743200]]></ColumnWidth><Table>',
        )
    )
    assert "ColumnWidth" in editing_context["dimensions"], "未提取列宽编辑上下文"
    assert editing_context["nativeLayoutState"]["hiddenColumns"], "未提取原生隐藏列上下文"
    assert any("原生隐藏配置" in item for item in editing_context["policy"]), "缺少原生隐藏配置编辑提示"
    full_replace_xml = _minimal_workbook("<Extra>full replace</Extra>")
    candidate = fr_report_high_authority_agent_service._build_candidate_xml(
        source_xml,
        {"finalCptXml": full_replace_xml},
        [],
    )
    assert "full replace" in candidate

    patch_candidate = fr_report_high_authority_agent_service._build_candidate_xml(
        source_xml,
        {
            "operations": [
                {
                    "operationType": "xml_patch",
                    "payload": {
                        "patches": [
                            {
                                "action": "replace",
                                "selector": "cell:B6",
                                "newXml": '<C c="1" r="5"><O><![CDATA[黄曲霉]]></O><PrivilegeControl/><Expand/></C>',
                            }
                        ]
                    },
                }
            ]
        },
        [
            {
                "operationType": "xml_patch",
                "payload": {
                    "patches": [
                        {
                            "action": "replace",
                            "selector": "cell:B6",
                            "newXml": '<C c="1" r="5"><O><![CDATA[黄曲霉]]></O><PrivilegeControl/><Expand/></C>',
                        }
                    ]
                },
            }
        ],
    )
    assert "黄曲霉" in patch_candidate and "旧值" not in patch_candidate

    missing_cell_candidate = fr_report_high_authority_agent_service._build_candidate_xml(
        source_xml,
        {},
        [
            {
                "operationType": "xml_patch",
                "payload": {
                    "patches": [
                        {
                            "action": "replace",
                            "selector": "WorkBook > Report > cell:B5",
                            "newXml": '<C><O><![CDATA[yyyy年MM月dd日]]></O><PrivilegeControl/><Expand/></C>',
                        }
                    ]
                },
            }
        ],
    )
    assert 'c="1"' in missing_cell_candidate and 'r="4"' in missing_cell_candidate and "yyyy年MM月dd日" in missing_cell_candidate

    slice_result = fr_report_high_authority_agent_service._read_cpt_slice(source_xml, "cell:B6")
    assert slice_result["status"] == "success" and "旧值" in slice_result["xml"], "ReAct CPT 片段读取工具未命中单元格"

    _event_name, _event_payload, observation = await fr_report_high_authority_agent_service._execute_react_tool(
        tool_name="apply_cpt_patch",
        arguments={
            "patches": [
                {
                    "action": "replace",
                    "selector": "cell:B6",
                    "newXml": '<C c="1" r="5"><O><![CDATA[循环工具已修改]]></O><PrivilegeControl/><Expand/></C>',
                }
            ]
        },
        source_xml=source_xml,
        database_source_context={},
        attachment_context=[],
        loop_state={},
    )
    assert observation["status"] == "success", f"ReAct apply_cpt_patch 工具未生成候选 CPT：{observation}"
    assert not fr_report_high_authority_agent_service._should_write_immediately("隐藏第一列", "review")
    assert fr_report_high_authority_agent_service._should_write_immediately("直接写入 CPT", "review")
    draft = fr_report_high_authority_agent_service._build_operation_draft(
        assistant_message="已生成修改项",
        normalized_path="webroot/APP/reportlets/test.cpt",
        operations=[
            {
                "operationType": "xml_patch",
                "summary": "替换 B6",
                "riskLevel": "medium",
                "payload": {"patches": [{"action": "replace", "selector": "cell:B6", "newXml": "<C/>"}]},
            }
        ],
        candidate_xml=source_xml,
        warnings=[],
    )
    assert draft["status"] == "draft" and draft["safety"]["requiresApproval"], "默认确认草稿未正确生成"

    event = fr_report_high_authority_agent_service._sse("message_delta", {"content": "读取 CPT"})
    assert re.search(r"event: message_delta\ndata: ", event), "SSE 事件格式不正确"

    await _run_optional_sqlserver_selftest()
    print("FineReport 高权限 Agent 轻量自测通过")


if __name__ == "__main__":
    asyncio.run(main())
