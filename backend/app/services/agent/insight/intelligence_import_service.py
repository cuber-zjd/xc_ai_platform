import json
import re
from datetime import datetime
from io import BytesIO
from typing import Any

from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from openpyxl import load_workbook
from sqlmodel.ext.asyncio.session import AsyncSession

from app.core.llm_factory import LLMFactory
from app.core.logger import logger
from app.schemas.agent.insight.intelligence import InsightIntelligenceImportItem, InsightIntelligenceImportPreviewResponse
from app.services.agent.insight.dictionary_service import insight_dictionary_service


class InsightIntelligenceImportService:
    max_file_size = 12 * 1024 * 1024
    max_text_chars = 60000
    allowed_suffixes = {".docx", ".xlsx", ".xlsm"}

    async def preview(
        self,
        db: AsyncSession,
        *,
        file_name: str,
        content: bytes,
    ) -> InsightIntelligenceImportPreviewResponse:
        suffix = self._suffix(file_name)
        if suffix not in self.allowed_suffixes:
            raise ValueError("仅支持 .docx、.xlsx、.xlsm 文件；旧版 .doc/.xls 请先转换后上传")
        if len(content) > self.max_file_size:
            raise ValueError("文件过大，请控制在 12MB 以内")

        text, warnings = self._extract_text(file_name, content)
        text = text[: self.max_text_chars]
        categories = await insight_dictionary_service.list_categories(db)
        tags = await insight_dictionary_service.list_tags(db)
        items = await self._split_with_ai(text, categories=categories, tags=tags, file_name=file_name)
        if not items:
            items = self._fallback_split(text, categories=categories, tags=tags)
            warnings.append("AI 拆分未返回可用结果，已使用规则拆分预览")
        return InsightIntelligenceImportPreviewResponse(
            file_name=file_name,
            file_type=suffix.lstrip("."),
            extracted_text_length=len(text),
            items=items[:80],
            warnings=warnings,
        )

    def _extract_text(self, file_name: str, content: bytes) -> tuple[str, list[str]]:
        suffix = self._suffix(file_name)
        if suffix == ".docx":
            return self._extract_docx(content), []
        if suffix in {".xlsx", ".xlsm"}:
            return self._extract_xlsx(content), []
        return "", ["文件类型不支持"]

    def _extract_docx(self, content: bytes) -> str:
        doc = Document(BytesIO(content))
        parts: list[str] = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip().replace("\n", " ") for cell in row.cells if cell.text.strip()]
                if values:
                    parts.append(" | ".join(values))
        return "\n".join(parts)

    def _extract_xlsx(self, content: bytes) -> str:
        workbook = load_workbook(BytesIO(content), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in workbook.worksheets[:8]:
            parts.append(f"【Sheet】{sheet.title}")
            row_count = 0
            for row in sheet.iter_rows(values_only=True):
                values = [self._cell_text(value) for value in row]
                values = [value for value in values if value]
                if not values:
                    continue
                parts.append(" | ".join(values))
                row_count += 1
                if row_count >= 300:
                    parts.append("【提示】当前 Sheet 超过 300 行，后续行已截断")
                    break
        workbook.close()
        return "\n".join(parts)

    async def _split_with_ai(self, text: str, *, categories: list[Any], tags: list[Any], file_name: str) -> list[InsightIntelligenceImportItem]:
        if not text.strip():
            return []
        category_payload = [{"code": item.category_code, "name": item.category_name} for item in categories]
        tag_payload = [{"code": item.tag_code, "name": item.tag_name, "category": item.tag_type} for item in tags]
        prompt = (
            "你是香驰控股研发营销市场洞察平台的情报整理助手。请从上传文件中拆分出多条可独立保存的正式情报。"
            "不要死板按段落拆分，一个 Word/Excel 可能包含多条情报，也可能某几行共同构成一条情报。"
            "分类和标签必须从给定字典中选择；没有合适标签时放入 suggested_new_tags，不要自由生成 tag_codes。"
            "只输出 JSON，格式为 {\"items\":[...]}。每条包含 title、summary、content、subject_type、subject_name、"
            "intelligence_type、category_code、tag_codes、importance_level、publish_time、selection_reason、business_insight、suggested_new_tags。"
        )
        try:
            response = await LLMFactory.safe_invoke(
                [
                    SystemMessage(content=prompt),
                    HumanMessage(
                        content=json.dumps(
                            {
                                "file_name": file_name,
                                "categories": category_payload,
                                "tags": tag_payload,
                                "document_text": text,
                            },
                            ensure_ascii=False,
                        )
                    ),
                ],
                capability="complex-reasoning",
            )
            raw = self._strip_json_fence(str(response.content if hasattr(response, "content") else response))
            data = json.loads(raw)
            raw_items = data.get("items") if isinstance(data, dict) else []
            return [self._normalize_item(item, categories=categories, tags=tags) for item in raw_items if isinstance(item, dict)]
        except Exception as exc:
            logger.warning("Insight 情报上传 AI 拆分失败: {}", exc)
            return []

    def _fallback_split(self, text: str, *, categories: list[Any], tags: list[Any]) -> list[InsightIntelligenceImportItem]:
        blocks = [block.strip() for block in re.split(r"\n\s*\n|(?=^【Sheet】)", text, flags=re.MULTILINE) if block.strip()]
        if len(blocks) <= 1:
            blocks = [line.strip() for line in text.splitlines() if len(line.strip()) >= 12]
        result = []
        default_category = categories[0].category_code if categories else None
        default_tag = tags[0].tag_code if tags else None
        default_tag_name = tags[0].tag_name if tags else None
        for block in blocks[:50]:
            title = block.splitlines()[0].strip(" |")[:80]
            result.append(
                InsightIntelligenceImportItem(
                    title=title,
                    summary=block[:300],
                    content=block,
                    subject_type="custom",
                    intelligence_type="行业资讯",
                    category_code=default_category,
                    category_name=categories[0].category_name if categories else None,
                    tag_codes=[default_tag] if default_tag else [],
                    tag_names=[default_tag_name] if default_tag_name else [],
                    importance_level="medium",
                    selection_reason="文件导入内容包含可独立跟进的信息，建议作为正式情报留存。",
                    business_insight="可供销售、市场或研发结合业务场景进一步判断。",
                )
            )
        return result

    def _normalize_item(self, item: dict[str, Any], *, categories: list[Any], tags: list[Any]) -> InsightIntelligenceImportItem:
        category_map = {row.category_code: row.category_name for row in categories}
        tag_map = {row.tag_code: row for row in tags}
        tag_codes = [str(code).strip() for code in item.get("tag_codes", []) if str(code).strip() in tag_map] if isinstance(item.get("tag_codes"), list) else []
        category_code = str(item.get("category_code") or "").strip() or None
        if category_code not in category_map:
            category_code = tag_map[tag_codes[0]].tag_type if tag_codes else (categories[0].category_code if categories else None)
        publish_time = self._parse_datetime(item.get("publish_time"))
        return InsightIntelligenceImportItem(
            title=str(item.get("title") or "未命名情报").strip()[:500],
            summary=self._optional_text(item.get("summary")),
            content=self._optional_text(item.get("content")),
            subject_type=str(item.get("subject_type") or "custom")[:30],
            subject_name=self._optional_text(item.get("subject_name"), limit=200),
            intelligence_type=str(item.get("intelligence_type") or category_map.get(category_code or "", "行业资讯"))[:50],
            category_code=category_code,
            category_name=category_map.get(category_code or ""),
            tag_codes=tag_codes,
            tag_names=[tag_map[code].tag_name for code in tag_codes],
            suggested_new_tags=[str(value).strip() for value in item.get("suggested_new_tags", []) if str(value).strip()] if isinstance(item.get("suggested_new_tags"), list) else [],
            importance_level=str(item.get("importance_level") or "medium")[:20],
            publish_time=publish_time,
            selection_reason=self._optional_text(item.get("selection_reason"), limit=1000),
            business_insight=self._optional_text(item.get("business_insight"), limit=1000),
            source_title=self._optional_text(item.get("source_title"), limit=500),
            source_url=self._optional_text(item.get("source_url"), limit=1000),
        )

    def _cell_text(self, value: object) -> str:
        if value is None:
            return ""
        if isinstance(value, datetime):
            return value.strftime("%Y-%m-%d %H:%M")
        return str(value).strip()

    def _optional_text(self, value: object, *, limit: int | None = None) -> str | None:
        text = str(value or "").strip()
        if not text:
            return None
        return text[:limit] if limit else text

    def _parse_datetime(self, value: object) -> datetime | None:
        if isinstance(value, datetime):
            return value
        text = str(value or "").strip()
        if not text:
            return None
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d", "%Y/%m/%d"):
            try:
                return datetime.strptime(text[:19], fmt)
            except ValueError:
                continue
        return None

    def _strip_json_fence(self, text: str) -> str:
        value = text.strip()
        if value.startswith("```"):
            value = re.sub(r"^```(?:json)?", "", value, flags=re.IGNORECASE).strip()
            value = re.sub(r"```$", "", value).strip()
        return value

    def _suffix(self, file_name: str) -> str:
        name = file_name.lower().strip()
        return "." + name.rsplit(".", 1)[-1] if "." in name else ""


insight_intelligence_import_service = InsightIntelligenceImportService()
