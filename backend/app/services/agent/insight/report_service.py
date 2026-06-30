import json
import re
import zipfile
from collections import Counter, defaultdict
from datetime import datetime
from html import escape
from io import BytesIO
from pathlib import Path
from collections.abc import Awaitable, Callable
from typing import Any
from uuid import uuid4
from xml.etree import ElementTree as ET

from langchain_core.messages import HumanMessage, SystemMessage
from sqlalchemy import exists, func, or_
from sqlmodel import select
from sqlmodel.ext.asyncio.session import AsyncSession

from app.core.llm_factory import LLMFactory
from app.models.agent.insight import (
    InsightCompany,
    InsightGraphEdge,
    InsightGraphNode,
    InsightIntelligence,
    InsightIntelligenceAsset,
    InsightIntelligenceSource,
    InsightReport,
    InsightReportExport,
    InsightReportMaterial,
    InsightReportPreference,
    InsightReportTemplate,
    InsightReportVersion,
    InsightTask,
    InsightTaskStatus,
)
from app.schemas.agent.insight.asset import InsightAssetSearchRequest
from app.schemas.agent.insight.report import (
    InsightReportChartPoint,
    InsightReportChartRead,
    InsightReportDetail,
    InsightReportExportRead,
    InsightReportGenerateRequest,
    InsightReportGenerateResponse,
    InsightReportListItem,
    InsightReportMaterialRead,
    InsightReportPreferenceRead,
    InsightReportPreferenceUpdate,
    InsightReportRead,
    InsightReportTemplateCreate,
    InsightReportTemplateCloneRequest,
    InsightReportTemplatePublishRequest,
    InsightReportTemplateRead,
    InsightReportTemplateSection,
    InsightReportTemplateUploadResponse,
    InsightReportTemplateUpdate,
    InsightReportUpdateRequest,
    InsightReportVersionRead,
)
from app.schemas.page import Page
from app.services.agent.insight.asset_service import insight_asset_service
from app.services.agent.insight.permission_service import insight_permission_service

ReportProgressCallback = Callable[[dict[str, Any]], Awaitable[None]]


REPORT_TEMPLATES: list[InsightReportTemplateRead] = [
    InsightReportTemplateRead(
        template_code="customer_business_review",
        template_name="客户经营洞察报告",
        description="面向客户成功、研发营销和销售跟进，围绕客户经营动态、机会、风险与合作建议形成正式研究报告。",
        report_type="专题报告",
        default_prompt="生成一份可直接交付的 Word 式客户经营洞察报告，正文完整、结论克制、引用可追溯，立场是帮助客户识别机会和风险。",
        sections=[
            InsightReportTemplateSection(section_key="background", heading="一、研究背景与客户概况", description="说明报告范围、证据口径、客户所处市场与当前关注重点。"),
            InsightReportTemplateSection(section_key="signals", heading="二、关键动态与证据交叉验证", description="按主题整合多来源资讯，说明哪些信号较强、哪些仍需验证。"),
            InsightReportTemplateSection(section_key="implications", heading="三、经营含义与合作机会", description="把公开资讯转化为客户经营、产品方案、渠道或合作跟进启发。"),
            InsightReportTemplateSection(section_key="risks", heading="四、风险提醒与待验证事项", description="客观指出不确定性、潜在风险和后续补充数据口径。"),
            InsightReportTemplateSection(section_key="recommendations", heading="五、结论与行动建议", description="形成可执行的客户维护、方案匹配和后续跟进建议。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="company_dynamic_report",
        template_name="企业动态跟进报告",
        description="聚焦单一企业或少数企业的新品、渠道、品牌、资本、供应链与风险动态。",
        report_type="企业动态报告",
        default_prompt="围绕目标企业的最新公开动态生成正式跟进报告，突出事实、影响、待验证点和下一步跟进动作。",
        sections=[
            InsightReportTemplateSection(section_key="overview", heading="一、企业近期动态总览", description="概括企业近阶段主要公开事件和信息密度。"),
            InsightReportTemplateSection(section_key="themes", heading="二、重点主题分析", description="围绕新品、渠道、供应链、品牌传播、组织动作等主题展开。"),
            InsightReportTemplateSection(section_key="impact", heading="三、对客户经营与合作的影响", description="说明这些动态对客户经营判断和合作机会的含义。"),
            InsightReportTemplateSection(section_key="watchlist", heading="四、后续观察清单", description="列出需要持续跟踪的事件、来源和验证方向。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="industry_topic_report",
        template_name="行业专题研究报告",
        description="用于行业、市场、产品、政策或技术主题研究，适合跨企业、多来源的专题判断。",
        report_type="专题报告",
        default_prompt="围绕行业专题生成正式研究报告，要求多来源归纳、证据分层、观点克制，并给出对客户经营和产品方案的启发。",
        sections=[
            InsightReportTemplateSection(section_key="topic_scope", heading="一、专题范围与研究口径", description="界定主题边界、时间范围和证据来源。"),
            InsightReportTemplateSection(section_key="market_signals", heading="二、市场信号与结构变化", description="归纳行业变化、需求趋势、竞争动作与政策影响。"),
            InsightReportTemplateSection(section_key="client_relevance", heading="三、客户相关性与业务启发", description="连接到目标客户、产品方案或合作场景。"),
            InsightReportTemplateSection(section_key="uncertainty", heading="四、不确定性与补充验证", description="说明证据缺口、冲突信息和后续数据需求。"),
            InsightReportTemplateSection(section_key="summary", heading="五、综合判断", description="输出谨慎但有方向感的专题结论。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="weekly_report",
        template_name="市场洞察周报",
        description="按周汇总客户、竞品、行业、政策与风险信号，适合例会和阶段跟进。",
        report_type="周报",
        default_prompt="生成一份市场洞察周报，要求按主题归纳本周高价值信号、客户影响、机会风险和下周关注事项。",
        sections=[
            InsightReportTemplateSection(section_key="weekly_summary", heading="一、本周核心摘要", description="用正式报告语气概括本周最重要的市场和客户信号。"),
            InsightReportTemplateSection(section_key="customer_updates", heading="二、客户与重点企业动态", description="汇总客户和重点企业的高价值公开动态。"),
            InsightReportTemplateSection(section_key="market_updates", heading="三、行业与市场变化", description="归纳行业趋势、竞品动作、政策与技术变化。"),
            InsightReportTemplateSection(section_key="next_week", heading="四、下周跟进建议", description="形成可执行的监测和业务跟进清单。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="competitor_dynamic_report",
        template_name="竞对动态报告",
        description="围绕竞对企业的新品、产能、渠道、专利、价格和风险动作形成可跟踪报告。",
        report_type="竞对报告",
        default_prompt="生成一份竞对动态报告，要求按竞对动作、业务影响、机会风险和后续监测重点组织，所有判断必须基于引用证据。",
        sections=[
            InsightReportTemplateSection(section_key="summary", heading="一、竞对动态摘要", description="概括本期高价值竞对信号和结论边界。"),
            InsightReportTemplateSection(section_key="moves", heading="二、重点竞对动作", description="按企业或主题梳理新品、渠道、产能、技术、资本等动作。"),
            InsightReportTemplateSection(section_key="impact", heading="三、对我方业务的影响", description="分析对研发、销售、客户维护和产品策略的影响。"),
            InsightReportTemplateSection(section_key="watchlist", heading="四、后续监测清单", description="列出需要继续跟踪的企业、关键词、来源和验证事项。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="customer_new_product_opportunity",
        template_name="客户新品机会报告",
        description="聚焦客户新品、配方、渠道和消费趋势，辅助研发营销识别合作切入点。",
        report_type="机会报告",
        default_prompt="生成一份客户新品机会报告，把公开新品与渠道动态转化为产品方案、样品推荐和客户拜访建议。",
        sections=[
            InsightReportTemplateSection(section_key="signals", heading="一、新品与渠道信号", description="汇总客户和下游品牌的新品、渠道、传播和消费趋势。"),
            InsightReportTemplateSection(section_key="needs", heading="二、潜在需求判断", description="结合证据判断甜味、功能、风味、成本、合规等潜在需求。"),
            InsightReportTemplateSection(section_key="opportunities", heading="三、合作机会与方案建议", description="提出可跟进客户、产品方向和销售动作。"),
            InsightReportTemplateSection(section_key="risks", heading="四、风险与待验证问题", description="说明需要业务或研发进一步验证的假设。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="rd_topic_trend_report",
        template_name="研发课题趋势报告",
        description="面向研发课题、技术路线、专利和配方趋势，沉淀可验证的研发方向。",
        report_type="研发报告",
        default_prompt="生成一份研发课题趋势报告，突出技术路线、专利信号、产品应用场景、机会风险和待验证实验方向。",
        sections=[
            InsightReportTemplateSection(section_key="topic_scope", heading="一、课题范围与证据口径", description="说明课题边界、证据来源和证据强弱。"),
            InsightReportTemplateSection(section_key="trend", heading="二、技术与应用趋势", description="归纳技术、专利、配方、应用场景和客户需求变化。"),
            InsightReportTemplateSection(section_key="rd_actions", heading="三、研发启发与验证建议", description="输出可落地的研发验证、样品开发或情报补采方向。"),
            InsightReportTemplateSection(section_key="uncertainty", heading="四、不确定性与补充数据", description="列出证据缺口和需要补充的实验或市场数据。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="policy_regulation_brief",
        template_name="政策/法规简报",
        description="汇总政策、法规、标准、监管与政府项目动态，适合法务、质量和市场同步。",
        report_type="政策简报",
        default_prompt="生成一份政策/法规简报，要求区分正式政策、征求意见、行业标准和政府项目信息，并说明影响范围。",
        sections=[
            InsightReportTemplateSection(section_key="policy_summary", heading="一、政策法规摘要", description="概括本期政策、法规、标准和监管重点。"),
            InsightReportTemplateSection(section_key="impact_scope", heading="二、影响范围分析", description="说明影响到的产品、客户、区域、合规或质量管理事项。"),
            InsightReportTemplateSection(section_key="actions", heading="三、应对建议", description="形成业务、研发、质量和客户沟通建议。"),
            InsightReportTemplateSection(section_key="monitoring", heading="四、后续关注", description="列出仍需跟踪的政策节点和信息来源。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="ecommerce_new_product_monitor",
        template_name="电商新品监测报告",
        description="跟踪电商平台新品、价格、卖点、配料和消费者反馈，服务市场与产品判断。",
        report_type="电商监测",
        default_prompt="生成一份电商新品监测报告，聚焦新品卖点、价格带、配料信号、消费者反馈和可借鉴机会。",
        sections=[
            InsightReportTemplateSection(section_key="new_products", heading="一、电商新品概览", description="按品牌或品类汇总新品、卖点、价格和渠道信息。"),
            InsightReportTemplateSection(section_key="product_signals", heading="二、产品与配方信号", description="提炼甜味、功能、风味、健康化、成本和包装趋势。"),
            InsightReportTemplateSection(section_key="opportunities", heading="三、市场机会与竞品启发", description="输出对研发、销售和客户方案的启发。"),
            InsightReportTemplateSection(section_key="risks", heading="四、风险与持续监测", description="说明证据局限、疑似营销信息和后续监测关键词。"),
        ],
    ),
    InsightReportTemplateRead(
        template_code="deep_research_report",
        template_name="深度研究报告",
        description="针对开放式研究问题，基于库内情报形成证据矩阵、结论、机会风险和后续验证问题。",
        report_type="深度研究",
        default_prompt="生成一份深度研究报告，必须先给结论，再列证据矩阵、机会风险、反证和后续待验证问题，禁止脱离证据编造。",
        sections=[
            InsightReportTemplateSection(section_key="answer", heading="一、研究结论", description="直接回答研究问题并说明置信边界。"),
            InsightReportTemplateSection(section_key="evidence_matrix", heading="二、证据矩阵", description="按证据主题列出情报、来源、日期和支撑关系。"),
            InsightReportTemplateSection(section_key="opportunity_risk", heading="三、机会与风险", description="综合机会点、风险点和业务影响。"),
            InsightReportTemplateSection(section_key="next_questions", heading="四、后续待验证问题", description="列出需要继续采集、访谈或内部验证的问题。"),
        ],
    ),
]

for template in REPORT_TEMPLATES:
    template.scope = "market"
    template.market_status = "listed"
    template.market_category = "系统内置"
    template.template_kind = "html"
    template.export_formats = ["html", "pdf", "docx"]
    template.visibility_scope = "public"


class InsightReportService:
    export_storage_root = Path(__file__).resolve().parents[4] / "storage" / "insight_exports"

    async def get_preference(self, db: AsyncSession, *, user_id: int) -> InsightReportPreferenceRead:
        row = await self._get_preference_row(db, user_id)
        if row:
            return self._to_preference_read(row)
        return self._default_preference(user_id)

    async def update_preference(
        self,
        db: AsyncSession,
        payload: InsightReportPreferenceUpdate,
        *,
        user_id: int,
    ) -> InsightReportPreferenceRead:
        row = await self._get_preference_row(db, user_id)
        if not row:
            row = InsightReportPreference(user_id=user_id, create_by=str(user_id), update_by=str(user_id))
            db.add(row)
            await db.flush()
        data = payload.model_dump(exclude_unset=True)
        for key, value in data.items():
            setattr(row, key, value)
        row.update_time = datetime.now()
        row.update_by = str(user_id)
        await db.commit()
        await db.refresh(row)
        return self._to_preference_read(row)

    async def list_templates(self, db: AsyncSession, *, user_id: int, is_admin: bool) -> list[InsightReportTemplateRead]:
        filters = [InsightReportTemplate.is_deleted == 0, InsightReportTemplate.status == "active"]
        if not is_admin:
            visibility_filter = await insight_permission_service.visibility_filter_for_user(
                db,
                InsightReportTemplate,
                target_type="report_template",
                user_id=user_id,
                is_admin=is_admin,
            )
            filters.append(
                or_(
                    InsightReportTemplate.scope == "market",
                    InsightReportTemplate.market_status == "listed",
                    InsightReportTemplate.owner_user_id == user_id,
                    visibility_filter,
                )
            )
        rows = list(
            (
                await db.exec(
                    select(InsightReportTemplate)
                    .where(*filters)
                    .order_by(InsightReportTemplate.update_time.desc(), InsightReportTemplate.id.desc())
                )
            ).all()
        )
        return [*REPORT_TEMPLATES, *[self._to_template_read(row) for row in rows]]

    async def create_template(
        self,
        db: AsyncSession,
        payload: InsightReportTemplateCreate,
        *,
        user_id: int,
    ) -> InsightReportTemplateRead:
        sections = payload.sections or self._default_template_sections()
        row = InsightReportTemplate(
            template_code=f"custom_{user_id}_{uuid4().hex[:12]}",
            template_name=payload.template_name,
            description=payload.description,
            report_type=payload.report_type,
            default_prompt=payload.default_prompt,
            sections_json=[section.model_dump(mode="json") for section in sections],
            structure_json=payload.structure_json or {},
            template_kind=payload.template_kind,
            export_formats=payload.export_formats or self._default_export_formats(payload.template_kind),
            style_code=payload.style_code,
            scope="personal",
            owner_user_id=user_id,
            visibility_scope=payload.visibility_scope,
            create_by=str(user_id),
            update_by=str(user_id),
        )
        db.add(row)
        await db.commit()
        await db.refresh(row)
        return self._to_template_read(row)

    async def create_template_from_upload(
        self,
        db: AsyncSession,
        *,
        file_name: str,
        file_bytes: bytes,
        template_name: str | None,
        report_type: str,
        description: str | None,
        user_id: int,
    ) -> InsightReportTemplateUploadResponse:
        if not file_bytes:
            raise ValueError("模板文件为空")
        if len(file_bytes) > 10 * 1024 * 1024:
            raise ValueError("模板文件不能超过 10MB")

        suffix = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if suffix not in {"docx", "xlsx"}:
            raise ValueError("当前仅支持上传 docx 或 xlsx 模板")

        parsed = self._parse_template_file(file_name, file_bytes)
        parsed["export_boundary"] = {
            "source_file_type": suffix,
            "parse_supported": True,
            "templated_export_supported": False,
            "message": "当前仅使用上传模板解析章节、表格和字段结构；DOCX/XLSX 套版导出尚未接入。",
        }
        sections = self._sections_from_structure(parsed)
        default_prompt = self._prompt_from_structure(parsed, report_type)
        row = InsightReportTemplate(
            template_code=f"upload_{user_id}_{uuid4().hex[:12]}",
            template_name=template_name or self._template_name_from_file(file_name),
            description=description or parsed.get("summary") or "由上传模板解析生成，可继续人工调整结构和 Prompt。",
            report_type=report_type or "专题报告",
            default_prompt=default_prompt,
            sections_json=[section.model_dump(mode="json") for section in sections],
            structure_json=parsed,
            template_kind="document",
            export_formats=["html", "pdf", "docx"],
            source_file_name=file_name[:300],
            source_file_type=suffix,
            source_file_size=len(file_bytes),
            scope="personal",
            owner_user_id=user_id,
            visibility_scope="private",
            create_by=str(user_id),
            update_by=str(user_id),
        )
        db.add(row)
        await db.commit()
        await db.refresh(row)
        return InsightReportTemplateUploadResponse(
            template=self._to_template_read(row),
            parsed_structure=parsed,
            extracted_text_preview=self._short_text(parsed.get("text_preview"), 1200),
        )

    async def update_template(
        self,
        db: AsyncSession,
        template_id: int,
        payload: InsightReportTemplateUpdate,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportTemplateRead:
        row = await self._get_custom_template_by_id(db, template_id, user_id=user_id, is_admin=is_admin)
        data = payload.model_dump(exclude_unset=True)
        sections = data.pop("sections", None)
        for key, value in data.items():
            if value is not None:
                setattr(row, key, value)
        if sections is not None:
            row.sections_json = [section.model_dump(mode="json") for section in sections]
        if data.get("structure_json") is not None:
            row.structure_json = data["structure_json"] or {}
        row.update_time = datetime.now()
        row.update_by = str(user_id)
        await db.commit()
        await db.refresh(row)
        return self._to_template_read(row)

    async def publish_template(
        self,
        db: AsyncSession,
        template_id: int,
        payload: InsightReportTemplatePublishRequest,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportTemplateRead:
        row = await self._get_custom_template_by_id(db, template_id, user_id=user_id, is_admin=is_admin)
        row.scope = "market"
        row.market_status = "listed"
        row.market_category = payload.market_category
        row.market_description = payload.market_description
        row.visibility_scope = "public"
        row.published_at = datetime.now()
        row.published_by_user_id = user_id
        row.update_time = datetime.now()
        row.update_by = str(user_id)
        await db.commit()
        await db.refresh(row)
        return self._to_template_read(row)

    async def clone_template(
        self,
        db: AsyncSession,
        template_code: str,
        payload: InsightReportTemplateCloneRequest,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportTemplateRead:
        source = await self._get_template(db, template_code, user_id=user_id, is_admin=is_admin)
        if source.scope not in {"system", "market"} and source.owner_user_id != user_id:
            raise ValueError("只能复制系统模板、市场模板或自己的模板")
        row = InsightReportTemplate(
            template_code=f"clone_{user_id}_{uuid4().hex[:12]}",
            template_name=payload.template_name or f"{source.template_name} 副本",
            description=source.description,
            report_type=source.report_type,
            default_prompt=source.default_prompt,
            sections_json=[section.model_dump(mode="json") for section in source.sections],
            structure_json=source.structure_json or {},
            template_kind=source.template_kind,
            style_code=source.style_code,
            export_formats=source.export_formats or self._default_export_formats(source.template_kind),
            scope="personal",
            market_status="not_listed",
            cloned_from_template_id=source.id,
            owner_user_id=user_id,
            visibility_scope="private",
            create_by=str(user_id),
            update_by=str(user_id),
        )
        db.add(row)
        await db.commit()
        await db.refresh(row)
        return self._to_template_read(row)

    async def delete_template(
        self,
        db: AsyncSession,
        template_id: int,
        *,
        user_id: int,
        is_admin: bool,
    ) -> None:
        row = await self._get_custom_template_by_id(db, template_id, user_id=user_id, is_admin=is_admin)
        row.is_deleted = 1
        row.status = "deleted"
        row.update_time = datetime.now()
        row.update_by = str(user_id)
        await db.commit()

    async def _get_template(
        self,
        db: AsyncSession,
        template_code: str | None,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportTemplateRead:
        for template in REPORT_TEMPLATES:
            if template.template_code == template_code:
                return template
        if template_code:
            filters = [
                InsightReportTemplate.template_code == template_code,
                InsightReportTemplate.is_deleted == 0,
                InsightReportTemplate.status == "active",
            ]
            if not is_admin:
                visibility_filter = await insight_permission_service.visibility_filter_for_user(
                    db,
                    InsightReportTemplate,
                    target_type="report_template",
                    user_id=user_id,
                    is_admin=is_admin,
                )
                filters.append(
                    or_(
                        InsightReportTemplate.scope == "market",
                        InsightReportTemplate.market_status == "listed",
                        InsightReportTemplate.owner_user_id == user_id,
                        visibility_filter,
                    )
                )
            row = (await db.exec(select(InsightReportTemplate).where(*filters))).first()
            if row:
                return self._to_template_read(row)
        return REPORT_TEMPLATES[0]

    async def list_reports(
        self,
        db: AsyncSession,
        *,
        page: int,
        size: int,
        keyword: str | None,
        report_type: str | None,
        status: str | None,
        user_id: int,
        is_admin: bool,
    ) -> Page[InsightReportListItem]:
        page = max(page, 1)
        size = min(max(size, 1), 100)
        filters = [InsightReport.is_deleted == 0]
        if keyword:
            like_keyword = f"%{keyword.strip()}%"
            filters.append(or_(InsightReport.title.ilike(like_keyword), InsightReport.summary.ilike(like_keyword)))
        if report_type:
            filters.append(InsightReport.report_type == report_type)
        if status:
            filters.append(InsightReport.status == status)
        if not is_admin:
            filters.append(await self._report_company_isolation_filter(db, user_id=user_id, is_admin=is_admin))
        filters.append(
            await insight_permission_service.visibility_filter_for_user(
                db,
                InsightReport,
                target_type="report",
                user_id=user_id,
                is_admin=is_admin,
            )
        )

        total = (await db.exec(select(func.count()).select_from(InsightReport).where(*filters))).one()
        rows = list(
            (
                await db.exec(
                    select(InsightReport)
                    .where(*filters)
                    .order_by(InsightReport.update_time.desc(), InsightReport.id.desc())
                    .offset((page - 1) * size)
                    .limit(size)
                )
            ).all()
        )
        return Page.create(items=[self._to_report_list_item(row) for row in rows], total=total, page=page, size=size)

    async def get_report_detail(
        self,
        db: AsyncSession,
        report_id: int,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportDetail:
        report = await self._get_report(db, report_id, user_id=user_id, is_admin=is_admin, permission="view")
        materials = await self._list_materials(db, report_id)
        versions = await self._list_versions(db, report_id)
        charts = await self._build_report_charts(db, materials)
        return self._to_report_detail(report, materials, versions, charts)

    async def generate_report(
        self,
        db: AsyncSession,
        payload: InsightReportGenerateRequest,
        *,
        user_id: int,
        is_admin: bool,
        progress_callback: ReportProgressCallback | None = None,
    ) -> InsightReportGenerateResponse:
        async def publish(step: str, title: str, detail: str, progress: int, **extra: Any) -> None:
            if progress_callback:
                await progress_callback(
                    {
                        "event": "progress",
                        "step": step,
                        "title": title,
                        "detail": detail,
                        "progress": progress,
                        **extra,
                    }
                )

        await publish("understand", "理解研究问题", "正在确认报告类型、研究范围和重点问题。", 5)
        task = InsightTask(
            task_uid=f"insight_report_{uuid4().hex}",
            task_type="report_generate",
            status=InsightTaskStatus.RUNNING,
            progress=10,
            started_at=datetime.now(),
            input_payload=payload.model_dump(mode="json"),
            create_by=str(user_id),
            update_by=str(user_id),
        )
        db.add(task)
        await db.flush()

        try:
            preference = await self.get_preference(db, user_id=user_id)
            payload = self._apply_preference_to_payload(payload, preference)
            await publish("search", "查找参考素材", "正在从已入库素材中查找与主题相关的内容。", 16)
            intelligence_rows = await self._select_materials(db, payload, user_id=user_id, is_admin=is_admin)
            if not intelligence_rows:
                raise ValueError("没有检索到可用于报告的正式情报资产，请先扩大主题、企业或时间范围，或补充采集。")
            await publish(
                "screen",
                "整理可用素材",
                f"已找到 {len(intelligence_rows)} 条候选素材，正在去重、分组并判断相关性。",
                32,
                material_count=len(intelligence_rows),
            )

            source_map = await self._list_primary_sources(db, [row.id for row in intelligence_rows if row.id])
            company_map = await self._list_companies(db, [row.company_id for row in intelligence_rows if row.company_id])
            asset_map = await self._asset_map_for_intelligences(db, [row.id for row in intelligence_rows if row.id])
            material_payload = [self._material_payload(row, source_map.get(row.id or 0), company_map, asset_map.get(row.id or 0)) for row in intelligence_rows]
            graph_context = await self._graph_context_for_materials(db, material_payload, user_id=user_id, is_admin=is_admin)
            await publish(
                "link",
                "补充关联线索",
                "正在按企业、产品、市场动作、风险和机会整理线索。",
                46,
                material_count=len(material_payload),
                relation_count=len(graph_context.get("edges", [])),
            )
            template = await self._get_template(db, payload.template_code, user_id=user_id, is_admin=is_admin)
            await publish("outline", "形成报告大纲", f"正在按“{template.template_name}”组织章节和重点观点。", 58)
            await publish("write", "撰写报告正文", "正在把素材整理成可直接查看的正式报告，并保留关键引用来源。", 70)
            content_json, generation_mode = await self._generate_content(payload, material_payload, template, graph_context)
            await publish("check", "检查报告质量", "正在检查空话、弱结论、重复内容和需要继续验证的问题。", 86)
            content_json["template_code"] = template.template_code
            content_json["template_name"] = template.template_name
            content_json["generation_mode"] = generation_mode
            content_json["evidence_retrieval"] = {
                "source": "formal_asset_rag",
                "query": await self._asset_query_from_payload(db, payload),
                "asset_count": len([item for item in material_payload if item.get("asset_id")]),
                "graph_node_count": len(graph_context.get("nodes", [])),
                "graph_edge_count": len(graph_context.get("edges", [])),
            }
            title = payload.title or content_json.get("title") or self._default_title(payload, material_payload)
            summary = self._short_text(content_json.get("executive_summary") or content_json.get("summary"), 1200)
            primary_company = self._primary_company(material_payload)

            report = InsightReport(
                report_uid=f"report_{uuid4().hex}",
                title=title,
                report_type=payload.report_type,
                period_start=payload.period_start,
                period_end=payload.period_end,
                company_id=primary_company.get("company_id") if primary_company else None,
                company_name=primary_company.get("company_name") if primary_company else None,
                content_json=content_json,
                summary=summary,
                status="final",
                version_no=1,
                material_count=len(material_payload),
                owner_user_id=user_id,
                visibility_scope="private",
                create_by=str(user_id),
                update_by=str(user_id),
            )
            db.add(report)
            await db.flush()

            for index, row in enumerate(material_payload, start=1):
                db.add(
                    InsightReportMaterial(
                        report_id=report.id or 0,
                        intelligence_id=row["id"],
                        section_key=self._section_key(row),
                        sort_no=index,
                        quote_text=row.get("summary"),
                        source_url=row.get("source_url"),
                        source_title=row.get("source_title"),
                        selection_source=self._selection_source(payload),
                        selection_reason=row.get("selection_reason"),
                    )
                )

            db.add(
                InsightReportVersion(
                    report_id=report.id or 0,
                    version_no=1,
                    content_json=content_json,
                    change_summary="首次生成正式报告",
                    created_by_user_id=user_id,
                )
            )
            task.report_id = report.id
            task.status = InsightTaskStatus.SUCCESS
            task.progress = 100
            task.finished_at = datetime.now()
            task.output_payload = {
                "report_id": report.id,
                "material_count": len(material_payload),
                "generation_mode": generation_mode,
            }
            await db.commit()
            await db.refresh(report)
            detail = await self.get_report_detail(db, report.id or 0, user_id=user_id, is_admin=True)
            await publish(
                "save",
                "保存正式报告",
                "报告已生成，正在切换到正文预览。",
                100,
                report_id=report.id,
                material_count=len(material_payload),
                generation_mode=generation_mode,
            )
            return InsightReportGenerateResponse(
                report=detail,
                task_id=task.id,
                used_material_count=len(material_payload),
                generation_mode=generation_mode,
            )
        except Exception as exc:
            task.status = InsightTaskStatus.FAILED
            task.progress = 100
            task.finished_at = datetime.now()
            task.error_message = str(exc)
            await db.commit()
            raise

    async def update_report(
        self,
        db: AsyncSession,
        report_id: int,
        payload: InsightReportUpdateRequest,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportDetail:
        report = await self._get_report(db, report_id, user_id=user_id, is_admin=is_admin, permission="edit")
        data = payload.model_dump(exclude_unset=True, exclude={"change_summary"})
        for key, value in data.items():
            if value is not None:
                setattr(report, key, value)
        report.version_no += 1
        report.update_time = datetime.now()
        report.update_by = str(user_id)
        db.add(
            InsightReportVersion(
                report_id=report.id or 0,
                version_no=report.version_no,
                content_json=report.content_json,
                change_summary=payload.change_summary or "人工更新报告草稿",
                created_by_user_id=user_id,
            )
        )
        await db.commit()
        await db.refresh(report)
        return await self.get_report_detail(db, report.id or 0, user_id=user_id, is_admin=is_admin)

    async def export_report(
        self,
        db: AsyncSession,
        report_id: int,
        *,
        export_format: str,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportExportRead:
        export_format = export_format.lower().strip()
        if export_format not in {"html", "pdf", "docx"}:
            raise ValueError("当前阶段仅支持导出 HTML、PDF 和 DOCX；XLSX 套版导出将在后续阶段接入")

        report = await self._get_report(db, report_id, user_id=user_id, is_admin=is_admin, permission="view")
        export = InsightReportExport(
            export_uid=f"report_export_{uuid4().hex}",
            report_id=report.id or 0,
            report_version_no=report.version_no,
            export_format=export_format,
            status="running",
            requested_by_user_id=user_id,
            storage_backend="local",
            create_by=str(user_id),
            update_by=str(user_id),
        )
        db.add(export)
        await db.commit()
        await db.refresh(export)

        try:
            detail = await self.get_report_detail(db, report.id or 0, user_id=user_id, is_admin=is_admin)
            html = self._render_report_html(detail)
            export_dir = self.export_storage_root / str(report.id)
            export_dir.mkdir(parents=True, exist_ok=True)
            file_name = self._export_file_name(report, export.id or 0, export_format)
            file_path = export_dir / file_name
            if export_format == "pdf":
                self._write_report_pdf(detail, file_path)
            elif export_format == "docx":
                self._write_report_docx(detail, file_path)
            else:
                file_path.write_text(html, encoding="utf-8")
            export.status = "success"
            export.file_name = file_name
            export.file_path = str(file_path)
            export.file_size = file_path.stat().st_size
            export.content_type = self._export_content_type(export_format)
            export.finished_at = datetime.now()
            export.update_time = datetime.now()
            export.update_by = str(user_id)
            await db.commit()
            await db.refresh(export)
            return self._to_export_read(export)
        except Exception as exc:
            export.status = "failed"
            export.error_message = str(exc)[:1000]
            export.finished_at = datetime.now()
            export.update_time = datetime.now()
            export.update_by = str(user_id)
            await db.commit()
            await db.refresh(export)
            return self._to_export_read(export)

    async def list_report_exports(
        self,
        db: AsyncSession,
        report_id: int,
        *,
        user_id: int,
        is_admin: bool,
    ) -> list[InsightReportExportRead]:
        report = await self._get_report(db, report_id, user_id=user_id, is_admin=is_admin, permission="view")
        rows = list(
            (
                await db.exec(
                    select(InsightReportExport)
                    .where(
                        InsightReportExport.report_id == (report.id or 0),
                        InsightReportExport.is_deleted == 0,
                    )
                    .order_by(InsightReportExport.create_time.desc(), InsightReportExport.id.desc())
                )
            ).all()
        )
        return [self._to_export_read(row) for row in rows]

    async def get_report_export_file(
        self,
        db: AsyncSession,
        report_id: int,
        export_id: int,
        *,
        user_id: int,
        is_admin: bool,
    ) -> tuple[Path, InsightReportExportRead]:
        report = await self._get_report(db, report_id, user_id=user_id, is_admin=is_admin, permission="view")
        export = (
            await db.exec(
                select(InsightReportExport).where(
                    InsightReportExport.id == export_id,
                    InsightReportExport.report_id == (report.id or 0),
                    InsightReportExport.is_deleted == 0,
                )
            )
        ).first()
        if not export:
            raise ValueError("导出记录不存在或无权访问")
        if export.status != "success" or not export.file_path:
            raise ValueError("导出文件尚未生成成功")
        file_path = Path(export.file_path)
        if not file_path.exists() or not file_path.is_file():
            raise ValueError("导出文件不存在或已被清理")
        return file_path, self._to_export_read(export)

    async def _select_materials(
        self,
        db: AsyncSession,
        payload: InsightReportGenerateRequest,
        *,
        user_id: int,
        is_admin: bool,
    ) -> list[InsightIntelligence]:
        filters = [InsightIntelligence.is_deleted == 0, InsightIntelligence.status == "active"]
        if not payload.intelligence_ids:
            filters.append(self._exclude_test_intelligence_filter())
        if payload.company_ids:
            filters.append(InsightIntelligence.company_id.in_(payload.company_ids))
        if payload.data_source_ids:
            filters.append(InsightIntelligence.data_source_id.in_(payload.data_source_ids))
        if payload.period_start:
            filters.append(or_(InsightIntelligence.publish_time >= payload.period_start, InsightIntelligence.create_time >= payload.period_start))
        if payload.period_end:
            filters.append(or_(InsightIntelligence.publish_time <= payload.period_end, InsightIntelligence.create_time <= payload.period_end))
        if not is_admin:
            filters.append(await self._intelligence_company_isolation_filter(db, user_id=user_id, is_admin=is_admin))
            filters.append(
                await insight_permission_service.visibility_filter_for_user(
                    db,
                    InsightIntelligence,
                    target_type="intelligence",
                    user_id=user_id,
                    is_admin=is_admin,
                )
            )

        if payload.intelligence_ids:
            filters.append(InsightIntelligence.id.in_(payload.intelligence_ids))
            statement = select(InsightIntelligence).where(*filters).limit(payload.max_materials)
            return list((await db.exec(statement)).all())

        if payload.data_source_ids:
            statement = (
                select(InsightIntelligence)
                .where(*filters)
                .order_by(InsightIntelligence.importance_level.desc(), InsightIntelligence.publish_time.desc().nullslast(), InsightIntelligence.create_time.desc())
                .limit(payload.max_materials)
            )
            return list((await db.exec(statement)).all())

        return await self._select_materials_from_assets(db, payload, filters, user_id=user_id, is_admin=is_admin)

    async def _select_materials_from_assets(
        self,
        db: AsyncSession,
        payload: InsightReportGenerateRequest,
        intelligence_filters: list[Any],
        *,
        user_id: int,
        is_admin: bool,
    ) -> list[InsightIntelligence]:
        query = await self._asset_query_from_payload(db, payload)
        selected_ids: list[int] = []
        try:
            search_result = await insight_asset_service.search_assets(
                db,
                InsightAssetSearchRequest(
                    query=query,
                    top_k=min(max(payload.max_materials, 5), 30),
                    include_candidates=False,
                    company_id=payload.company_ids[0] if len(payload.company_ids) == 1 else None,
                    date_from=payload.period_start,
                    date_to=payload.period_end,
                ),
                user_id=user_id,
                is_admin=is_admin,
            )
            for hit in search_result.hits:
                intelligence_id = hit.asset.intelligence_id
                if intelligence_id and intelligence_id not in selected_ids:
                    selected_ids.append(intelligence_id)
        except Exception:
            selected_ids = []

        if selected_ids:
            rows = list(
                (
                    await db.exec(
                        select(InsightIntelligence)
                        .where(*intelligence_filters, InsightIntelligence.id.in_(selected_ids))
                        .limit(payload.max_materials)
                    )
                ).all()
            )
            row_map = {row.id: row for row in rows}
            ordered = [row_map[item_id] for item_id in selected_ids if item_id in row_map]
            if ordered:
                return ordered[: payload.max_materials]
        else:
            ordered = []

        direct_filters = list(intelligence_filters)
        if selected_ids:
            direct_filters.append(InsightIntelligence.id.not_in(selected_ids))
        relevance_filter = self._keyword_relevance_filter(query)
        if relevance_filter is not None:
            direct_filters.append(relevance_filter)
        direct_rows = list(
            (
                await db.exec(
                    select(InsightIntelligence)
                    .where(*direct_filters)
                    .order_by(InsightIntelligence.importance_level.desc(), InsightIntelligence.publish_time.desc().nullslast(), InsightIntelligence.create_time.desc())
                    .limit(max(payload.max_materials - len(ordered), 0))
                )
            ).all()
        )
        return (ordered + direct_rows)[: payload.max_materials]

    def _exclude_test_intelligence_filter(self) -> Any:
        test_terms = ("测试", "烟测", "smoke", "test", "demo", "样本", "P1企业档案测试", "测试客户")
        clauses = []
        for term in test_terms:
            like_term = f"%{term}%"
            clauses.extend(
                [
                    InsightIntelligence.title.ilike(like_term),
                    InsightIntelligence.summary.ilike(like_term),
                    InsightIntelligence.subject_name.ilike(like_term),
                ]
            )
        return ~or_(*clauses)

    def _keyword_relevance_filter(self, query: str) -> Any | None:
        stop_words = {
            "专题报告",
            "深度研究",
            "客户经营洞察",
            "正式素材库",
            "报告模板",
            "模板用途",
            "用户重点问题",
            "研究范围",
            "全部可用素材",
            "香驰控股",
            "机会",
            "风险",
            "战略",
            "客户",
            "竞对",
        }
        terms: list[str] = []
        for raw_term in re.split(r"[\s,，。；;：:\n\r、/|]+", query or ""):
            term = raw_term.strip()
            if len(term) < 2 or term in stop_words or term in terms:
                continue
            terms.append(term[:80])
            if len(terms) >= 12:
                break
        if not terms:
            return None
        clauses = []
        for term in terms:
            like_term = f"%{term}%"
            clauses.extend(
                [
                    InsightIntelligence.title.ilike(like_term),
                    InsightIntelligence.summary.ilike(like_term),
                    InsightIntelligence.content.ilike(like_term),
                    InsightIntelligence.subject_name.ilike(like_term),
                    InsightIntelligence.business_domain.ilike(like_term),
                ]
            )
        return or_(*clauses)

    async def _asset_query_from_payload(self, db: AsyncSession, payload: InsightReportGenerateRequest) -> str:
        parts = [
            payload.title,
            payload.report_type,
            payload.generation_prompt,
            "香驰控股 大豆 玉米 果葡糖浆 麦芽糖 植物蛋白 豆粕 粮油 客户 竞对 机会 风险 战略",
        ]
        if payload.company_ids:
            rows = list(
                (
                    await db.exec(
                        select(InsightCompany).where(
                            InsightCompany.id.in_(payload.company_ids),
                            InsightCompany.is_deleted == 0,
                        )
                    )
                ).all()
            )
            for row in rows:
                parts.extend([row.name, row.short_name, row.company_type, row.industry])
        return " ".join(str(part).strip() for part in parts if str(part or "").strip())[:1000]

    async def _asset_map_for_intelligences(
        self,
        db: AsyncSession,
        intelligence_ids: list[int],
    ) -> dict[int, InsightIntelligenceAsset]:
        if not intelligence_ids:
            return {}
        rows = list(
            (
                await db.exec(
                    select(InsightIntelligenceAsset).where(
                        InsightIntelligenceAsset.intelligence_id.in_(intelligence_ids),
                        InsightIntelligenceAsset.is_deleted == 0,
                    )
                )
            ).all()
        )
        return {row.intelligence_id: row for row in rows if row.intelligence_id}

    async def _graph_context_for_materials(
        self,
        db: AsyncSession,
        materials: list[dict[str, Any]],
        *,
        user_id: int,
        is_admin: bool,
    ) -> dict[str, Any]:
        asset_ids = [int(item["asset_id"]) for item in materials if item.get("asset_id")]
        if not asset_ids:
            return {"nodes": [], "edges": []}
        visible_asset_filters = [
            InsightIntelligenceAsset.id.in_(asset_ids),
            InsightIntelligenceAsset.is_deleted == 0,
            InsightIntelligenceAsset.status == "active",
        ]
        if not is_admin:
            visible_asset_filters.append(
                await insight_permission_service.visibility_filter_for_user(
                    db,
                    InsightIntelligenceAsset,
                    target_type="asset",
                    user_id=user_id,
                    is_admin=is_admin,
                )
            )
        visible_asset_ids = list((await db.exec(select(InsightIntelligenceAsset.id).where(*visible_asset_filters))).all())
        if not visible_asset_ids:
            return {"nodes": [], "edges": []}
        edges = list(
            (
                await db.exec(
                    select(InsightGraphEdge)
                    .where(
                        InsightGraphEdge.source_asset_id.in_(visible_asset_ids),
                        InsightGraphEdge.is_deleted == 0,
                        InsightGraphEdge.status == "active",
                    )
                    .order_by(InsightGraphEdge.confidence.desc(), InsightGraphEdge.update_time.desc())
                    .limit(80)
                )
            ).all()
        )
        node_ids = sorted({edge.source_node_id for edge in edges} | {edge.target_node_id for edge in edges})
        nodes = list(
            (
                await db.exec(
                    select(InsightGraphNode)
                    .where(
                        InsightGraphNode.id.in_(node_ids) if node_ids else InsightGraphNode.source_asset_id.in_(visible_asset_ids),
                        InsightGraphNode.is_deleted == 0,
                        InsightGraphNode.status == "active",
                    )
                    .limit(80)
                )
            ).all()
        )
        node_map = {node.id: node for node in nodes}
        return {
            "nodes": [
                {
                    "id": node.id,
                    "type": node.node_type,
                    "name": node.node_name,
                    "company_id": node.company_id,
                }
                for node in nodes
            ],
            "edges": [
                {
                    "source": node_map.get(edge.source_node_id).node_name if node_map.get(edge.source_node_id) else edge.source_node_id,
                    "target": node_map.get(edge.target_node_id).node_name if node_map.get(edge.target_node_id) else edge.target_node_id,
                    "relation": edge.relation_type,
                    "confidence": edge.confidence,
                    "asset_id": edge.source_asset_id,
                    "evidence": self._short_text(edge.evidence_text, 260),
                }
                for edge in edges
            ],
        }

    async def _generate_content(
        self,
        payload: InsightReportGenerateRequest,
        materials: list[dict[str, Any]],
        template: InsightReportTemplateRead,
        graph_context: dict[str, Any] | None = None,
    ) -> tuple[dict[str, Any], str]:
        fallback = self._fallback_content(payload, materials, template)
        try:
            research_notes = await self._generate_research_notes(payload, materials, template, graph_context or {})
            messages = [
                SystemMessage(
                    content=(
                        "你是研发营销市场洞察平台的报告生成 Agent。"
                        "只能基于给定参考素材写报告，不得编造事实。"
                        "默认报告对象是我们的客户或潜在客户，报告立场是客户经营洞察和客户成功支持，不是第三方投研唱空、媒体评论或竞品攻击。"
                        "可以客观指出风险、压力和不确定性，但表达方式应服务于客户维护、销售跟进、产品方案匹配和合作机会识别。"
                        "涉及负面事件时避免情绪化措辞，优先写成'需要关注/需要验证/可能影响客户经营的信号'。"
                        "你将收到上一步研究整理出的研究问题、素材分组、强弱信号、缺口和建议大纲。"
                        "最终报告正文必须像 Word 研究报告，不要把'研究过程'、'思考过程'、'证据矩阵'作为正文大章节展示。"
                        "正文应包含摘要、背景、分章节分析、客户经营含义、合作机会、风险提醒、结论与建议，语气正式、段落完整、少用列表。"
                        "报告应优先保留对研发、市场、营销、销售有直接价值的新品、品类、消费、价格、渠道、客户痛点和合作机会。"
                        "最终只输出报告 JSON，不输出隐藏推理；reflection 字段用简短中文说明素材强弱、关联线索和仍需验证的问题。"
                        "素材中如包含 content_excerpt，表示系统已提供可参考的正文摘录；只有当摘要不足以支撑判断时，才结合正文摘录补充细节。"
                        "对供应链数字化、环保、资本市场、雇主品牌等偏远信息，除非用户明确要求或与销售切入直接相关，否则只作为背景，不要占据主要篇幅。"
                        "如果不同素材互相矛盾，要在正文中自然说明冲突和不确定性。"
                        "输出严格 JSON，字段包括 title、executive_summary、chapters、conclusion、key_findings、company_sections、risks、opportunities、evidence_matrix、reflection、follow_up_questions、source_notes。"
                        "chapters 是数组，每项包含 heading、paragraphs、evidence_ids；paragraphs 是适合直接放进报告正文的中文段落。"
                        "正文 paragraphs 不得出现'素材ID'、'证据ID'、'证据编号'、'ID:'等内部编号表达；引用关系只放在 evidence_ids 字段，由报告渲染层展示。"
                        "必须优先按照用户选择的报告模板组织 chapters，章节标题尽量与模板章节一致；如果证据不足，可以合并相邻章节，但不能输出空洞模板话。"
                        "所有标题、类型、原因和正文必须使用中文。"
                    )
                ),
                HumanMessage(
                    content=json.dumps(
                        {
                            "report_type": payload.report_type,
                            "template": template.model_dump(mode="json"),
                            "material_scope": {
                                "company_ids": payload.company_ids,
                                "data_source_ids": payload.data_source_ids,
                                "material_source": "正式素材库",
                                "max_materials": payload.max_materials,
                            },
                            "user_prompt": payload.generation_prompt,
                            "research_notes": research_notes,
                            "materials": materials[:120],
                        },
                        ensure_ascii=False,
                    )
                ),
            ]
            response = await LLMFactory.safe_invoke(
                messages,
                capability="complex-reasoning",
                temperature=0.25,
                json_mode=True,
                enable_reasoning=True,
                max_retries=4,
            )
            content = self._parse_llm_json(getattr(response, "content", str(response)))
            normalized = self._normalize_content(content, fallback)
            normalized["research_method"] = normalized.get("research_method") or [
                "先理解研究问题并拆成可回答的子问题。",
                "整理已入库素材，筛掉重复和明显无关的内容。",
                "按企业、产品、风险和机会归类，形成报告大纲后再撰写正文。",
                "检查结论是否有素材支撑，并标出仍需验证的问题。",
            ]
            normalized["research_process"] = research_notes
            return normalized, "asset_rag_deep_research"
        except Exception:
            return fallback, "rules"

    async def _generate_research_notes(
        self,
        payload: InsightReportGenerateRequest,
        materials: list[dict[str, Any]],
        template: InsightReportTemplateRead,
        graph_context: dict[str, Any],
    ) -> dict[str, Any]:
        fallback = self._fallback_research_notes(payload, materials, template, graph_context)
        try:
            messages = [
                SystemMessage(
                    content=(
                        "你是企业市场洞察研究员。请先做研究整理，不写正式报告。"
                        "你需要把用户问题拆成子问题，按企业、产品、市场、政策、风险和机会整理素材，"
                        "判断哪些素材支撑强、哪些只是弱线索，哪些地方缺资料。"
                        "只输出 JSON，字段包括 research_questions、material_groups、strong_signals、weak_signals、gaps、outline、quality_checks。"
                        "所有内容必须使用中文，避免技术术语。"
                    )
                ),
                HumanMessage(
                    content=json.dumps(
                        {
                            "report_type": payload.report_type,
                            "template_name": template.template_name,
                            "template_sections": [section.model_dump(mode="json") for section in template.sections],
                            "user_prompt": payload.generation_prompt,
                            "materials": materials[:120],
                            "relations": graph_context,
                        },
                        ensure_ascii=False,
                    )
                ),
            ]
            response = await LLMFactory.safe_invoke(
                messages,
                capability="complex-reasoning",
                temperature=0.15,
                json_mode=True,
                enable_reasoning=True,
                max_retries=2,
            )
            notes = self._parse_llm_json(getattr(response, "content", str(response)))
            return notes if isinstance(notes, dict) else fallback
        except Exception:
            return fallback

    def _fallback_research_notes(
        self,
        payload: InsightReportGenerateRequest,
        materials: list[dict[str, Any]],
        template: InsightReportTemplateRead,
        graph_context: dict[str, Any],
    ) -> dict[str, Any]:
        type_counts = Counter(item.get("intelligence_type") or "未分类" for item in materials)
        company_counts = Counter(item.get("company_name") or item.get("subject_name") or "未归属主题" for item in materials)
        groups = [
            {
                "name": name,
                "material_count": count,
                "focus": "优先检查是否与用户问题、客户机会或风险判断直接相关。",
            }
            for name, count in type_counts.most_common(8)
        ]
        return {
            "research_questions": [
                payload.generation_prompt or f"围绕{payload.report_type}形成可交付报告",
                "哪些变化对客户经营、销售跟进、产品方案或风险预警最有价值？",
                "哪些判断素材充分，哪些仍需要后续验证？",
            ],
            "material_groups": groups,
            "strong_signals": [
                {
                    "title": f"{name}相关素材较多",
                    "summary": f"当前共有 {count} 条相关素材，可作为报告重点方向之一。",
                }
                for name, count in company_counts.most_common(5)
                if count >= 2
            ],
            "weak_signals": [
                {
                    "title": item.get("title"),
                    "summary": item.get("summary") or "该素材可作为线索，但仍需要更多来源验证。",
                }
                for item in materials[:5]
            ],
            "gaps": [
                "当前主要依赖已入库公开素材，部分主题可能缺少企业官网、公告、招投标、招聘或内部经营数据印证。",
                "如果报告用于重要决策，建议后续补充关键企业的官方来源和内部客户合作记录。",
            ],
            "outline": [section.heading for section in template.sections] or ["摘要", "重点分析", "机会风险", "结论建议"],
            "quality_checks": [
                "优先写有素材支撑的判断。",
                "弱线索只作为待验证问题，不写成确定结论。",
                "避免空泛口号，尽量给出可跟进动作。",
                f"已识别 {len(graph_context.get('edges', [])) if isinstance(graph_context, dict) else 0} 条关联线索。",
            ],
        }

    def _fallback_content(
        self,
        payload: InsightReportGenerateRequest,
        materials: list[dict[str, Any]],
        template: InsightReportTemplateRead,
    ) -> dict[str, Any]:
        type_counts = Counter(item.get("intelligence_type") or "未分类" for item in materials)
        company_groups: dict[str, list[dict[str, Any]]] = defaultdict(list)
        for item in materials:
            company_groups[item.get("company_name") or item.get("subject_name") or "未归属"].append(item)
        key_findings = [
            {
                "title": item["title"],
                "insight": item.get("summary") or "该情报可作为后续研判线索。",
                "evidence_ids": [item["id"]],
            }
            for item in materials[:8]
        ]
        evidence_matrix = [
            {
                "theme": theme,
                "material_count": count,
                "evidence_strength": "中",
                "note": f"该主题目前有 {count} 条正式素材支撑，需结合后续来源继续验证。",
            }
            for theme, count in type_counts.most_common(8)
        ]
        company_sections = [
            {
                "company_name": company_name,
                "summary": f"共引用 {len(rows)} 条素材，重点集中在 {self._top_type_text(rows)}。",
                "signals": [{"title": row["title"], "summary": row.get("summary"), "evidence_ids": [row["id"]]} for row in rows[:6]],
            }
            for company_name, rows in company_groups.items()
        ]
        title = payload.title or self._default_title(payload, materials)
        lead_items = [self._material_sentence(item) for item in materials[:5]]
        lead_text = "；".join(item for item in lead_items if item)
        return {
            "title": title,
            "template_code": template.template_code,
            "template_name": template.template_name,
            "executive_summary": (
                f"本报告基于 {len(materials)} 条正式素材整理，覆盖 {len(company_groups)} 个企业或主题。"
                f"当前可直接使用的素材包括：{lead_text or '暂无可概括的核心素材'}。"
                f"类型分布以 {', '.join([f'{k}{v}条' for k, v in type_counts.most_common(5)]) or '未分类素材'} 为主。"
            ),
            "chapters": self._fallback_chapters(company_sections, key_findings, materials, template),
            "conclusion": "总体看，当前素材适合形成阶段性整理稿。涉及客户经营判断、销售跟进优先级和合作机会判断时，仍需要结合企业官方公告、渠道数据和内部客户合作记录继续验证。",
            "research_method": [
                "按企业和主题聚合正式素材，先识别高频议题和异常信号。",
                "对同一企业下的新品、经营、市场、风险信息进行交叉比对，避免单条资讯直接下结论。",
                "将结论分为已被多条素材支撑的发现、需要持续观察的机会和素材不足的风险假设。",
            ],
            "evidence_matrix": evidence_matrix,
            "key_findings": key_findings,
            "company_sections": company_sections,
            "risks": self._risk_items(materials),
            "opportunities": self._opportunity_items(materials),
            "reflection": [
                "当前报告主要依赖公开资讯和已入库网页正文，暂未接入工商变更、招投标、招聘和内部客户合作数据。",
                "部分证据缺少文章发布时间或完整正文，相关结论只能作为弱信号处理。",
                "同一事件可能被多个媒体转载，已做 URL 和正文去重，但仍需要在正式报告前人工复核关键事实。",
            ],
            "follow_up_questions": [
                "哪些企业的新品或市场动作出现连续多源印证，值得转入重点跟进？",
                "哪些风险只来自单一来源，需要补充官网、公告或第三方企业数据验证？",
                "后续接入启信宝、招聘和内部合作数据后，哪些结论可能被增强或推翻？",
            ],
            "source_notes": [
                "报告仅引用已入库正式情报和来源素材。",
                "竞品、行业和供应链背景素材会作为上下文进入报告，但不等同于目标企业自身事件。",
            ],
            "stats": {
                "material_count": len(materials),
                "type_counts": dict(type_counts.most_common(12)),
                "company_counts": {key: len(value) for key, value in company_groups.items()},
            },
        }

    def _risk_items(self, materials: list[dict[str, Any]]) -> list[dict[str, Any]]:
        keywords = ("下滑", "亏损", "价格战", "监管", "风险", "停售", "闭店", "竞争")
        return [
            {"title": item["title"], "summary": item.get("summary"), "evidence_ids": [item["id"]]}
            for item in materials
            if any(keyword in f"{item.get('title')} {item.get('summary')}" for keyword in keywords)
        ][:8]

    def _opportunity_items(self, materials: list[dict[str, Any]]) -> list[dict[str, Any]]:
        keywords = ("新品", "合作", "增长", "扩张", "出海", "上市", "投资", "供应链", "科研")
        return [
            {"title": item["title"], "summary": item.get("summary"), "evidence_ids": [item["id"]]}
            for item in materials
            if any(keyword in f"{item.get('title')} {item.get('summary')}" for keyword in keywords)
        ][:8]

    def _fallback_chapters(
        self,
        company_sections: list[dict[str, Any]],
        key_findings: list[dict[str, Any]],
        materials: list[dict[str, Any]],
        template: InsightReportTemplateRead | None = None,
    ) -> list[dict[str, Any]]:
        template_sections = template.sections if template else []
        company_paragraphs = [
            f"{section['company_name']}相关素材显示，{section['summary']}主要素材包括："
            + "；".join(self._material_sentence(signal) for signal in section.get("signals", [])[:3] if self._material_sentence(signal))
            + "。"
            for section in company_sections[:4]
            if section.get("summary")
        ]
        finding_paragraphs = [
            f"{item['title']}：{item.get('insight') or item.get('summary')}"
            for item in key_findings[:5]
            if item.get("title")
        ]
        return [
            {
                "heading": template_sections[0].heading if len(template_sections) > 0 else "一、市场与企业动态概览",
                "paragraphs": company_paragraphs
                or [f"本期共引用 {len(materials)} 条正式素材，覆盖目标企业动态、行业趋势、产品创新和潜在风险。"],
                "evidence_ids": [item["id"] for item in materials[:6] if item.get("id")],
            },
            {
                "heading": template_sections[1].heading if len(template_sections) > 1 else "二、关键发现与业务含义",
                "paragraphs": finding_paragraphs
                or ["现有素材显示，目标企业的公开动作主要集中在新品、渠道、供应链和品牌传播等方向，后续需要结合更多结构化数据验证其持续性。"],
                "evidence_ids": [item["id"] for item in materials[6:12] if item.get("id")],
            },
            {
                "heading": template_sections[2].heading if len(template_sections) > 2 else "三、后续跟踪建议",
                "paragraphs": [
                    "建议后续将公开资讯与企业工商变化、招投标、招聘、渠道价格和内部客户合作记录进行联动分析，以便把单点资讯转化为更稳定的客户经营判断，并为客户维护、产品方案匹配和合作机会跟进提供依据。",
                ],
                "evidence_ids": [item["id"] for item in materials[12:18] if item.get("id")],
            },
        ]

    def _material_sentence(self, item: dict[str, Any]) -> str:
        title = self._short_text(item.get("title") or "", 120)
        summary = self._short_text(item.get("summary") or item.get("insight") or "", 180)
        if title and summary:
            return f"{title}，{summary}"
        return title or summary

    def _normalize_content(self, content: dict[str, Any], fallback: dict[str, Any]) -> dict[str, Any]:
        result = fallback | {key: value for key, value in content.items() if value}
        for key in (
            "chapters",
            "research_method",
            "evidence_matrix",
            "key_findings",
            "company_sections",
            "risks",
            "opportunities",
            "reflection",
            "follow_up_questions",
            "source_notes",
        ):
            if not isinstance(result.get(key), list):
                result[key] = fallback.get(key, [])
        if not isinstance(result.get("conclusion"), str):
            result["conclusion"] = fallback.get("conclusion")
        if not isinstance(result.get("stats"), dict):
            result["stats"] = fallback["stats"]
        return result

    def _parse_llm_json(self, text: str) -> dict[str, Any]:
        value = text.strip()
        if value.startswith("```"):
            value = re.sub(r"^```(?:json)?", "", value, flags=re.IGNORECASE).strip()
            value = re.sub(r"```$", "", value).strip()
        parsed = json.loads(value)
        return parsed if isinstance(parsed, dict) else {}

    async def _get_preference_row(self, db: AsyncSession, user_id: int) -> InsightReportPreference | None:
        return (
            await db.exec(
                select(InsightReportPreference).where(
                    InsightReportPreference.user_id == user_id,
                    InsightReportPreference.is_deleted == 0,
                    InsightReportPreference.status == "active",
                )
            )
        ).first()

    def _default_preference(self, user_id: int) -> InsightReportPreferenceRead:
        now = datetime.now()
        return InsightReportPreferenceRead(
            id=0,
            create_time=now,
            update_time=now,
            user_id=user_id,
            default_template_code="customer_business_review",
            default_report_type="专题报告",
            default_folder_name=None,
            default_max_materials=100,
            writing_stance="客户经营视角",
            report_depth="深度研究",
            citation_style="正文上标引用",
            include_risks=True,
            include_opportunities=True,
            include_follow_up_questions=True,
            custom_prompt_suffix=None,
            status="active",
        )

    def _to_preference_read(self, row: InsightReportPreference) -> InsightReportPreferenceRead:
        return InsightReportPreferenceRead(
            id=row.id or 0,
            create_time=row.create_time,
            update_time=row.update_time,
            user_id=row.user_id,
            default_template_code=row.default_template_code,
            default_report_type=row.default_report_type,
            default_folder_name=row.default_folder_name,
            default_max_materials=row.default_max_materials,
            writing_stance=row.writing_stance,
            report_depth=row.report_depth,
            citation_style=row.citation_style,
            include_risks=row.include_risks,
            include_opportunities=row.include_opportunities,
            include_follow_up_questions=row.include_follow_up_questions,
            custom_prompt_suffix=row.custom_prompt_suffix,
            status=row.status,
        )

    def _apply_preference_to_payload(
        self,
        payload: InsightReportGenerateRequest,
        preference: InsightReportPreferenceRead,
    ) -> InsightReportGenerateRequest:
        data = payload.model_dump()
        if not data.get("template_code"):
            data["template_code"] = preference.default_template_code
        if not data.get("report_type"):
            data["report_type"] = preference.default_report_type
        if not data.get("folder_name"):
            data["folder_name"] = preference.default_folder_name
        if not data.get("max_materials"):
            data["max_materials"] = preference.default_max_materials
        prompt_parts = [
            data.get("generation_prompt"),
            f"写作立场：{preference.writing_stance}。",
            f"报告深度：{preference.report_depth}。",
            f"引用方式：{preference.citation_style}。",
            "需要包含风险提醒。" if preference.include_risks else "除非证据强相关，否则减少风险提醒篇幅。",
            "需要包含合作机会和业务建议。" if preference.include_opportunities else "减少合作机会展开，优先事实归纳。",
            "需要包含后续跟进问题。" if preference.include_follow_up_questions else "不需要单独输出后续跟进问题。",
            preference.custom_prompt_suffix,
        ]
        data["generation_prompt"] = "\n".join(part for part in prompt_parts if part)
        return InsightReportGenerateRequest(**data)

    def _selection_source(self, payload: InsightReportGenerateRequest) -> str:
        if payload.intelligence_ids:
            return "manual_ids"
        if payload.data_source_ids:
            return "data_source_filter"
        if payload.company_ids:
            return "company_asset_rag"
        return "formal_asset_rag"

    def _parse_template_file(self, file_name: str, file_bytes: bytes) -> dict[str, Any]:
        suffix = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else ""
        if suffix == "docx":
            return self._parse_docx_template(file_name, file_bytes)
        if suffix == "xlsx":
            return self._parse_xlsx_template(file_name, file_bytes)
        raise ValueError("当前仅支持上传 docx 或 xlsx 模板")

    def _parse_docx_template(self, file_name: str, file_bytes: bytes) -> dict[str, Any]:
        try:
            with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
                xml = archive.read("word/document.xml")
        except Exception as exc:
            raise ValueError("Word 模板解析失败，请确认文件为有效 docx") from exc

        root = ET.fromstring(xml)
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        paragraphs: list[dict[str, Any]] = []
        tables: list[dict[str, Any]] = []
        text_parts: list[str] = []

        for block in root.findall(".//w:body/*", ns):
            tag = block.tag.rsplit("}", 1)[-1]
            if tag == "p":
                text = "".join(node.text or "" for node in block.findall(".//w:t", ns)).strip()
                if not text:
                    continue
                style = block.find(".//w:pStyle", ns)
                style_value = style.attrib.get(f"{{{ns['w']}}}val") if style is not None else ""
                level = self._heading_level(style_value, text)
                paragraphs.append({"text": text, "style": style_value, "heading_level": level})
                text_parts.append(text)
            elif tag == "tbl":
                rows: list[list[str]] = []
                for tr in block.findall(".//w:tr", ns):
                    cells = []
                    for tc in tr.findall("./w:tc", ns):
                        cells.append("".join(node.text or "" for node in tc.findall(".//w:t", ns)).strip())
                    if any(cells):
                        rows.append(cells)
                if rows:
                    tables.append({"row_count": len(rows), "column_count": max(len(row) for row in rows), "sample_rows": rows[:5]})

        headings = [item for item in paragraphs if item["heading_level"]]
        return {
            "source": "upload",
            "file_type": "docx",
            "file_name": file_name,
            "paragraph_count": len(paragraphs),
            "table_count": len(tables),
            "headings": headings[:30],
            "tables": tables[:10],
            "text_preview": "\n".join(text_parts[:80]),
            "summary": f"Word 模板解析到 {len(headings)} 个标题、{len(paragraphs)} 个段落、{len(tables)} 个表格。",
        }

    def _parse_xlsx_template(self, file_name: str, file_bytes: bytes) -> dict[str, Any]:
        try:
            with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
                shared_strings = self._xlsx_shared_strings(archive)
                workbook = ET.fromstring(archive.read("xl/workbook.xml"))
                sheet_names = [node.attrib.get("name", "未命名Sheet") for node in workbook.findall(".//{*}sheet")]
                worksheet_names = sorted(name for name in archive.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml"))
                sheets = []
                for index, worksheet_name in enumerate(worksheet_names[:12]):
                    sheet_xml = ET.fromstring(archive.read(worksheet_name))
                    rows = self._xlsx_rows(sheet_xml, shared_strings)
                    non_empty_rows = [row for row in rows if any(cell for cell in row)]
                    sheets.append(
                        {
                            "sheet_name": sheet_names[index] if index < len(sheet_names) else f"Sheet{index + 1}",
                            "row_count": len(non_empty_rows),
                            "column_count": max((len(row) for row in non_empty_rows), default=0),
                            "sample_rows": non_empty_rows[:8],
                            "possible_headers": self._guess_xlsx_headers(non_empty_rows),
                        }
                    )
        except Exception as exc:
            raise ValueError("Excel 模板解析失败，请确认文件为有效 xlsx") from exc

        text_preview = []
        for sheet in sheets:
            text_preview.append(f"[{sheet['sheet_name']}]")
            for row in sheet["sample_rows"][:5]:
                text_preview.append(" | ".join(row))
        return {
            "source": "upload",
            "file_type": "xlsx",
            "file_name": file_name,
            "sheet_count": len(sheets),
            "sheets": sheets,
            "text_preview": "\n".join(text_preview),
            "summary": f"Excel 模板解析到 {len(sheets)} 个工作表，可按表格标题、字段顺序和样例行约束报告输出。",
        }

    def _xlsx_shared_strings(self, archive: zipfile.ZipFile) -> list[str]:
        if "xl/sharedStrings.xml" not in archive.namelist():
            return []
        root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
        return ["".join(node.text or "" for node in item.findall(".//{*}t")) for item in root.findall(".//{*}si")]

    def _xlsx_rows(self, root: ET.Element, shared_strings: list[str]) -> list[list[str]]:
        result: list[list[str]] = []
        for row in root.findall(".//{*}sheetData/{*}row")[:80]:
            cells: list[str] = []
            for cell in row.findall("{*}c"):
                cell_type = cell.attrib.get("t")
                value_node = cell.find("{*}v")
                inline_node = cell.find(".//{*}t")
                value = ""
                if cell_type == "s" and value_node is not None:
                    index = int(value_node.text or 0)
                    value = shared_strings[index] if 0 <= index < len(shared_strings) else ""
                elif inline_node is not None:
                    value = inline_node.text or ""
                elif value_node is not None:
                    value = value_node.text or ""
                cells.append(value.strip())
            result.append(cells)
        return result

    def _guess_xlsx_headers(self, rows: list[list[str]]) -> list[str]:
        for row in rows[:8]:
            useful = [cell for cell in row if cell.strip()]
            if len(useful) >= 2:
                return useful[:20]
        return []

    def _heading_level(self, style_value: str | None, text: str) -> int | None:
        style = (style_value or "").lower()
        if "heading1" in style or "1" == style[-1:] or re.match(r"^[一二三四五六七八九十]+[、.．]", text):
            return 1
        if "heading2" in style or re.match(r"^（?[一二三四五六七八九十]+[）)]", text):
            return 2
        if "heading3" in style or re.match(r"^\d+[.．、]", text):
            return 3
        if len(text) <= 30 and any(keyword in text for keyword in ("摘要", "背景", "分析", "结论", "建议", "风险", "机会")):
            return 2
        return None

    def _sections_from_structure(self, structure: dict[str, Any]) -> list[InsightReportTemplateSection]:
        if structure.get("file_type") == "docx":
            headings = structure.get("headings") if isinstance(structure.get("headings"), list) else []
            sections = [
                InsightReportTemplateSection(
                    section_key=f"section_{index}",
                    heading=str(item.get("text") or f"第 {index} 节")[:120],
                    description="参考上传 Word 模板中的标题层级和上下文，生成正式报告正文。",
                )
                for index, item in enumerate(headings, start=1)
                if isinstance(item, dict) and item.get("heading_level") in {1, 2}
            ]
            if sections:
                return sections[:12]
        if structure.get("file_type") == "xlsx":
            sheets = structure.get("sheets") if isinstance(structure.get("sheets"), list) else []
            sections = []
            for index, sheet in enumerate(sheets, start=1):
                if not isinstance(sheet, dict):
                    continue
                headers = "、".join(str(item) for item in (sheet.get("possible_headers") or [])[:8])
                sections.append(
                    InsightReportTemplateSection(
                        section_key=f"sheet_{index}",
                        heading=str(sheet.get("sheet_name") or f"Sheet{index}")[:120],
                        description=f"参考该 Sheet 的字段与表格结构组织内容。主要字段：{headers or '未识别到明确表头'}。",
                    )
                )
            if sections:
                return sections[:12]
        return self._default_template_sections()

    def _prompt_from_structure(self, structure: dict[str, Any], report_type: str) -> str:
        base = [
            f"请按上传的 {structure.get('file_type', '').upper()} 模板结构生成{report_type or '专题报告'}。",
            "报告正文必须服务于客户经营洞察，基于已选情报证据写作，不编造事实。",
            "尽量保留模板的标题层级、字段顺序、表格逻辑和措辞风格；模板中像占位符、表格字段、Sheet 名称的内容，应转化为报告输出约束。",
            "如果证据不足以填满某个模板模块，需要明确说明证据不足，不要空泛补字。",
        ]
        if structure.get("summary"):
            base.append(str(structure["summary"]))
        return "\n".join(base)

    def _template_name_from_file(self, file_name: str) -> str:
        stem = file_name.rsplit(".", 1)[0]
        return f"{stem[:80]} 模板"

    def _default_export_formats(self, template_kind: str | None) -> list[str]:
        _ = template_kind
        return ["html", "pdf", "docx"]

    async def _get_custom_template_by_id(
        self,
        db: AsyncSession,
        template_id: int,
        *,
        user_id: int,
        is_admin: bool,
    ) -> InsightReportTemplate:
        filters = [
            InsightReportTemplate.id == template_id,
            InsightReportTemplate.is_deleted == 0,
        ]
        if not is_admin:
            filters.append(
                await insight_permission_service.visibility_filter_for_user(
                    db,
                    InsightReportTemplate,
                    target_type="report_template",
                    user_id=user_id,
                    is_admin=is_admin,
                    permission="edit",
                )
            )
        row = (await db.exec(select(InsightReportTemplate).where(*filters))).first()
        if not row:
            raise ValueError("报告模板不存在或无权访问")
        return row

    def _default_template_sections(self) -> list[InsightReportTemplateSection]:
        return [
            InsightReportTemplateSection(section_key="summary", heading="一、核心摘要", description="概括报告主要发现和判断边界。"),
            InsightReportTemplateSection(section_key="analysis", heading="二、重点分析", description="围绕证据做正式正文分析。"),
            InsightReportTemplateSection(section_key="recommendations", heading="三、结论与建议", description="形成客户经营和业务跟进建议。"),
        ]

    def _to_template_read(self, row: InsightReportTemplate) -> InsightReportTemplateRead:
        sections = [
            InsightReportTemplateSection.model_validate(section)
            for section in (row.sections_json or [])
            if isinstance(section, dict)
        ]
        return InsightReportTemplateRead(
            id=row.id,
            template_code=row.template_code,
            template_name=row.template_name,
            description=row.description or "",
            report_type=row.report_type,
            default_prompt=row.default_prompt or "",
            sections=sections,
            structure_json=row.structure_json or {},
            template_kind=row.template_kind,
            style_code=row.style_code,
            export_formats=row.export_formats if row.export_formats is not None else self._default_export_formats(row.template_kind),
            source_file_name=row.source_file_name,
            source_file_type=row.source_file_type,
            source_file_size=row.source_file_size,
            scope=row.scope,
            market_status=row.market_status,
            market_category=row.market_category,
            market_description=row.market_description,
            cloned_from_template_id=row.cloned_from_template_id,
            published_at=row.published_at,
            published_by_user_id=row.published_by_user_id,
            owner_user_id=row.owner_user_id,
            owner_dept_id=row.owner_dept_id,
            visibility_scope=row.visibility_scope,
            editable=True,
        )

    async def _get_report(
        self,
        db: AsyncSession,
        report_id: int,
        *,
        user_id: int,
        is_admin: bool,
        permission: str = "view",
    ) -> InsightReport:
        filters = [InsightReport.id == report_id, InsightReport.is_deleted == 0]
        if not is_admin:
            filters.append(await self._report_company_isolation_filter(db, user_id=user_id, is_admin=is_admin))
        filters.append(
            await insight_permission_service.visibility_filter_for_user(
                db,
                InsightReport,
                target_type="report",
                user_id=user_id,
                is_admin=is_admin,
                permission=permission,
            )
        )
        report = (await db.exec(select(InsightReport).where(*filters))).first()
        if not report:
            raise ValueError("报告不存在或无权访问")
        return report

    async def _report_company_isolation_filter(
        self,
        db: AsyncSession,
        *,
        user_id: int | None,
        is_admin: bool,
    ):
        if is_admin:
            return True
        sys_company_id = await insight_permission_service.resolve_user_sys_company_id(db, user_id)
        if sys_company_id is None:
            return InsightReport.company_id.is_(None)
        return or_(
            InsightReport.company_id.is_(None),
            exists()
            .where(InsightCompany.id == InsightReport.company_id)
            .where(InsightCompany.sys_company_id == sys_company_id)
            .where(InsightCompany.is_deleted == 0),
        )

    async def _intelligence_company_isolation_filter(
        self,
        db: AsyncSession,
        *,
        user_id: int | None,
        is_admin: bool,
    ):
        if is_admin:
            return True
        sys_company_id = await insight_permission_service.resolve_user_sys_company_id(db, user_id)
        if sys_company_id is None:
            return InsightIntelligence.company_id.is_(None)
        return or_(
            InsightIntelligence.company_id.is_(None),
            exists()
            .where(InsightCompany.id == InsightIntelligence.company_id)
            .where(InsightCompany.sys_company_id == sys_company_id)
            .where(InsightCompany.is_deleted == 0),
        )

    async def _list_primary_sources(
        self,
        db: AsyncSession,
        intelligence_ids: list[int],
    ) -> dict[int, InsightIntelligenceSource]:
        if not intelligence_ids:
            return {}
        sources = list(
            (
                await db.exec(
                    select(InsightIntelligenceSource)
                    .where(
                        InsightIntelligenceSource.intelligence_id.in_(intelligence_ids),
                        InsightIntelligenceSource.is_deleted == 0,
                    )
                    .order_by(InsightIntelligenceSource.intelligence_id.asc(), InsightIntelligenceSource.id.asc())
                )
            ).all()
        )
        result: dict[int, InsightIntelligenceSource] = {}
        for source in sources:
            result.setdefault(source.intelligence_id, source)
        return result

    async def _list_companies(self, db: AsyncSession, company_ids: list[int]) -> dict[int, InsightCompany]:
        if not company_ids:
            return {}
        companies = list((await db.exec(select(InsightCompany).where(InsightCompany.id.in_(company_ids)))).all())
        return {company.id or 0: company for company in companies}

    async def _list_materials(self, db: AsyncSession, report_id: int) -> list[InsightReportMaterialRead]:
        materials = list(
            (
                await db.exec(
                    select(InsightReportMaterial)
                    .where(InsightReportMaterial.report_id == report_id, InsightReportMaterial.is_deleted == 0)
                    .order_by(InsightReportMaterial.sort_no.asc(), InsightReportMaterial.id.asc())
                )
            ).all()
        )
        intelligence_ids = [item.intelligence_id for item in materials]
        intelligences = list((await db.exec(select(InsightIntelligence).where(InsightIntelligence.id.in_(intelligence_ids)))).all()) if intelligence_ids else []
        intelligence_map = {item.id or 0: item for item in intelligences}
        return [self._to_material_read(item, intelligence_map.get(item.intelligence_id)) for item in materials]

    async def _list_versions(self, db: AsyncSession, report_id: int) -> list[InsightReportVersionRead]:
        rows = list(
            (
                await db.exec(
                    select(InsightReportVersion)
                    .where(InsightReportVersion.report_id == report_id, InsightReportVersion.is_deleted == 0)
                    .order_by(InsightReportVersion.version_no.desc(), InsightReportVersion.id.desc())
                )
            ).all()
        )
        return [self._to_version_read(row) for row in rows]

    async def _build_report_charts(
        self,
        db: AsyncSession,
        materials: list[InsightReportMaterialRead],
    ) -> list[InsightReportChartRead]:
        intelligence_ids = [item.intelligence_id for item in materials if item.intelligence_id]
        if not intelligence_ids:
            return []

        rows = list((await db.exec(select(InsightIntelligence).where(InsightIntelligence.id.in_(intelligence_ids)))).all())
        source_map = await self._list_primary_sources(db, intelligence_ids)
        company_map = await self._list_companies(db, [row.company_id for row in rows if row.company_id])

        subject_counts: Counter[str] = Counter()
        type_counts: Counter[str] = Counter()
        source_counts: Counter[str] = Counter()
        tag_counts: Counter[str] = Counter()
        signal_counts: Counter[str] = Counter()
        trend_counts: Counter[str] = Counter()

        for row in rows:
            company = company_map.get(row.company_id or 0)
            subject_name = (company.short_name or company.name) if company else row.subject_name
            subject_counts[subject_name or "未归属主题"] += 1
            type_counts[self._intelligence_type_label(row.intelligence_type)] += 1
            signal_counts[self._signal_bucket(row)] += 1

            source = source_map.get(row.id or 0)
            source_counts[self._source_type_label(source.source_type if source else None)] += 1

            event_time = row.publish_time or row.create_time
            if event_time:
                trend_counts[event_time.strftime("%m-%d")] += 1

            raw_payload = row.raw_payload or {}
            tags = raw_payload.get("suggested_tags") if isinstance(raw_payload, dict) else None
            if isinstance(tags, list):
                for tag in tags:
                    name = tag.get("name") if isinstance(tag, dict) else tag
                    if isinstance(name, str) and name.strip():
                        tag_counts[name.strip()] += 1

        charts = [
            self._counter_chart(
                "subject_distribution",
                "企业与主题证据分布",
                subject_counts,
                chart_type="bar",
                description="按报告引用证据归属的企业或主题统计。",
            ),
            self._counter_chart(
                "type_distribution",
                "情报类型分布",
                type_counts,
                chart_type="donut",
                description="观察报告证据主要集中在哪些经营信号。",
            ),
            self._counter_chart(
                "source_distribution",
                "来源渠道占比",
                source_counts,
                chart_type="donut",
                description="按主来源渠道统计，用于判断证据来源结构。",
            ),
            self._trend_chart(trend_counts),
            self._counter_chart(
                "signal_distribution",
                "机会与风险信号",
                signal_counts,
                chart_type="bar",
                description="基于标题、摘要和类型关键词的粗粒度信号划分。",
            ),
            self._counter_chart(
                "tag_frequency",
                "高频标签",
                tag_counts,
                chart_type="list",
                description="来自候选情报或正式情报中的标签聚合。",
            ),
        ]
        return [chart for chart in charts if chart.points]

    def _counter_chart(
        self,
        chart_key: str,
        title: str,
        counter: Counter[str],
        *,
        chart_type: str,
        description: str | None = None,
        limit: int = 8,
    ) -> InsightReportChartRead:
        total = sum(counter.values())
        points = [
            InsightReportChartPoint(
                key=str(label),
                label=str(label),
                value=count,
                percent=round(count * 100 / total, 1) if total else 0,
            )
            for label, count in counter.most_common(limit)
            if count > 0
        ]
        return InsightReportChartRead(
            chart_key=chart_key,
            title=title,
            description=description,
            chart_type=chart_type,
            unit="条",
            points=points,
        )

    def _trend_chart(self, counter: Counter[str]) -> InsightReportChartRead:
        points = [
            InsightReportChartPoint(key=label, label=label, value=count)
            for label, count in sorted(counter.items())[-12:]
            if count > 0
        ]
        return InsightReportChartRead(
            chart_key="publish_trend",
            title="证据发布时间趋势",
            description="按文章发布时间聚合；若来源未提供发布时间，则使用情报创建时间兜底。",
            chart_type="line",
            unit="条",
            points=points,
        )

    def _source_type_label(self, source_type: str | None) -> str:
        labels = {
            "baidu_news": "百度资讯",
            "bocha": "博查搜索",
            "bocha_search": "博查搜索",
            "bocha_news": "博查资讯",
            "bocha_web": "博查网页",
            "official_site": "官网",
            "web_page": "通用网页",
            "firecrawl": "网页抓取",
            "manual": "人工录入",
            "search": "搜索发现",
        }
        if not source_type:
            return "未知来源"
        return labels.get(source_type.lower(), source_type)

    def _intelligence_type_label(self, value: str | None) -> str:
        labels = {
            "market_trend": "市场趋势",
            "product_launch": "新品发布",
            "product_launch_failure": "新品受阻",
            "marketing_strategy": "营销策略",
            "strategic_planning": "战略规划",
            "competitor_strategy": "竞品策略",
            "channel_expansion": "渠道扩张",
            "supply_chain": "供应链动态",
            "risk_signal": "风险信号",
            "policy": "政策法规",
            "custom": "自定义",
        }
        if not value:
            return "未分类"
        return labels.get(value.lower(), value)

    def _signal_bucket(self, row: InsightIntelligence) -> str:
        text = f"{row.intelligence_type} {row.title} {row.summary or ''}".lower()
        if any(keyword in text for keyword in ("风险", "预警", "处罚", "下滑", "亏损", "停产", "召回", "舆情", "risk")):
            return "风险提醒"
        if any(keyword in text for keyword in ("机会", "合作", "新品", "增长", "扩张", "投资", "出海", "升级", "launch", "growth")):
            return "机会信号"
        return "常规动态"

    def _material_payload(
        self,
        row: InsightIntelligence,
        source: InsightIntelligenceSource | None,
        company_map: dict[int, InsightCompany],
        asset: InsightIntelligenceAsset | None = None,
    ) -> dict[str, Any]:
        company = company_map.get(row.company_id or 0)
        raw_payload = row.raw_payload or {}
        structured_payload = asset.structured_payload if asset and isinstance(asset.structured_payload, dict) else {}
        review_payload = asset.review_payload if asset and isinstance(asset.review_payload, dict) else raw_payload.get("ai_review")
        return {
            "id": row.id,
            "asset_id": asset.id if asset else None,
            "title": row.title,
            "summary": self._short_text(row.summary, 600),
            "company_id": row.company_id,
            "company_name": (company.short_name or company.name) if company else row.subject_name,
            "subject_name": row.subject_name,
            "intelligence_type": row.intelligence_type,
            "business_value": row.business_domain or (asset.business_value if asset else None),
            "importance_level": row.importance_level,
            "confidence": asset.confidence if asset else None,
            "publish_time": (row.publish_time or row.create_time).isoformat() if (row.publish_time or row.create_time) else None,
            "source_url": source.source_url if source else None,
            "source_title": source.source_title if source else None,
            "content_excerpt": self._short_text((asset.evidence_text if asset else None) or row.content, 1800) if self._should_include_content_excerpt(row) else None,
            "tags": (asset.tags if asset else raw_payload.get("suggested_tags")) if isinstance(raw_payload, dict) else [],
            "entities": asset.entities if asset else structured_payload.get("entities", []),
            "related_products": structured_payload.get("related_products", []),
            "opportunities": asset.opportunities if asset else structured_payload.get("opportunities", []),
            "risks": asset.risks if asset else structured_payload.get("risks", []),
            "keywords": asset.keywords if asset else [],
            "review_reason": review_payload.get("reason") if isinstance(review_payload, dict) else None,
            "selection_reason": f"{row.importance_level} 关注度，{row.intelligence_type} 类资产；{row.business_domain or (asset.business_value if asset else '')}".strip("；"),
        }

    def _should_include_content_excerpt(self, row: InsightIntelligence) -> bool:
        text = f"{row.intelligence_type} {row.title} {row.summary or ''} {row.content or ''}".lower()
        positive_keywords = (
            "新品",
            "上新",
            "口味",
            "果茶",
            "奶茶",
            "冰饮",
            "功能饮料",
            "无糖",
            "甜",
            "糖浆",
            "价格",
            "涨价",
            "促销",
            "门店",
            "私域",
            "会员",
            "下沉",
            "加盟",
            "海外",
            "出海",
            "供应地",
            "原料",
            "研发",
        )
        remote_keywords = (
            "数字化",
            "支付",
            "环保",
            "公益",
            "财报",
            "股价",
            "装瓶业务",
            "上市",
            "雇主品牌",
        )
        positive_score = sum(1 for keyword in positive_keywords if keyword in text)
        remote_score = sum(1 for keyword in remote_keywords if keyword in text)
        return positive_score > 0 and positive_score >= remote_score

    def _default_title(self, payload: InsightReportGenerateRequest, materials: list[dict[str, Any]]) -> str:
        primary = self._primary_company(materials)
        prefix = primary.get("company_name") if primary else "市场洞察"
        return f"{prefix}{payload.report_type}草稿"

    def _primary_company(self, materials: list[dict[str, Any]]) -> dict[str, Any] | None:
        rows = [item for item in materials if item.get("company_id")]
        if not rows:
            return None
        company_id, _ = Counter(item["company_id"] for item in rows).most_common(1)[0]
        match = next(item for item in rows if item["company_id"] == company_id)
        return {"company_id": company_id, "company_name": match.get("company_name")}

    def _section_key(self, row: dict[str, Any]) -> str:
        value = str(row.get("intelligence_type") or "key_findings").lower()
        if "risk" in value or "price" in value:
            return "risks"
        if "product" in value or "market" in value or "strategy" in value:
            return "opportunities"
        return "key_findings"

    def _top_type_text(self, rows: list[dict[str, Any]]) -> str:
        return "、".join(name for name, _count in Counter(row.get("intelligence_type") or "未分类" for row in rows).most_common(3))

    def _short_text(self, value: object, limit: int) -> str | None:
        if not isinstance(value, str):
            return None
        return value.strip()[:limit]

    def _to_report_list_item(self, row: InsightReport) -> InsightReportListItem:
        return InsightReportListItem(**self._to_report_read(row).model_dump())

    def _to_report_read(self, row: InsightReport) -> InsightReportRead:
        return InsightReportRead(
            id=row.id or 0,
            create_time=row.create_time,
            update_time=row.update_time,
            report_uid=row.report_uid,
            title=row.title,
            report_type=row.report_type,
            period_start=row.period_start,
            period_end=row.period_end,
            company_id=row.company_id,
            company_name=row.company_name,
            content_json=row.content_json or {},
            summary=row.summary,
            status=row.status,
            version_no=row.version_no,
            material_count=row.material_count,
            owner_user_id=row.owner_user_id,
            owner_dept_id=row.owner_dept_id,
            visibility_scope=row.visibility_scope,
        )

    def _to_export_read(self, row: InsightReportExport) -> InsightReportExportRead:
        return InsightReportExportRead(
            id=row.id or 0,
            create_time=row.create_time,
            update_time=row.update_time,
            export_uid=row.export_uid,
            report_id=row.report_id,
            report_version_no=row.report_version_no,
            export_format=row.export_format,
            status=row.status,
            file_name=row.file_name,
            file_size=row.file_size,
            content_type=row.content_type,
            storage_backend=row.storage_backend,
            error_message=row.error_message,
            requested_by_user_id=row.requested_by_user_id,
            finished_at=row.finished_at,
        )

    def _export_file_name(self, report: InsightReport, export_id: int, export_format: str = "html") -> str:
        title = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", report.title).strip("_")[:80] or "insight_report"
        suffix = export_format if export_format in {"html", "pdf", "docx"} else "html"
        return f"{title}_v{report.version_no}_export_{export_id}.{suffix}"

    def _export_content_type(self, export_format: str) -> str:
        if export_format == "pdf":
            return "application/pdf"
        if export_format == "docx":
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return "text/html; charset=utf-8"

    def _write_report_pdf(self, report: InsightReportDetail, file_path: Path) -> None:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_LEFT
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise RuntimeError("缺少 reportlab 依赖，无法生成 PDF") from exc

        font_name = self._register_pdf_font(pdfmetrics, TTFont)
        styles = getSampleStyleSheet()
        base = ParagraphStyle(
            "InsightBase",
            parent=styles["Normal"],
            fontName=font_name,
            fontSize=10.5,
            leading=18,
            textColor=colors.HexColor("#344846"),
            alignment=TA_LEFT,
            wordWrap="CJK",
            spaceAfter=8,
        )
        title_style = ParagraphStyle(
            "InsightTitle",
            parent=base,
            fontSize=22,
            leading=30,
            textColor=colors.HexColor("#10201f"),
            spaceAfter=12,
        )
        h2_style = ParagraphStyle(
            "InsightHeading",
            parent=base,
            fontSize=15,
            leading=22,
            textColor=colors.HexColor("#0f3f3a"),
            spaceBefore=14,
            spaceAfter=8,
        )
        meta_style = ParagraphStyle(
            "InsightMeta",
            parent=base,
            fontSize=9,
            leading=14,
            textColor=colors.HexColor("#667a77"),
        )
        small_style = ParagraphStyle(
            "InsightSmall",
            parent=base,
            fontSize=9,
            leading=14,
            textColor=colors.HexColor("#506461"),
        )
        content = report.content_json or {}
        chapters = content.get("chapters") if isinstance(content, dict) else []
        if not isinstance(chapters, list):
            chapters = []
        materials = report.materials or []
        material_map = {item.intelligence_id: item for item in materials}
        story: list[Any] = [
            Paragraph(self._pdf_escape(report.title), title_style),
            Paragraph(
                self._pdf_escape(
                    f"报告类型：{report.report_type} · 版本：第 {report.version_no} 版 · "
                    f"证据：{report.material_count} 条 · 导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
                ),
                meta_style,
            ),
            Spacer(1, 8),
            self._pdf_info_box(
                "摘要",
                self._strip_inline_evidence_ids(str(content.get("executive_summary") or report.summary or "暂无摘要。")),
                font_name,
                base,
                Table,
                TableStyle,
                colors,
            ),
        ]
        for chapter in chapters:
            if not isinstance(chapter, dict):
                continue
            story.append(Paragraph(self._pdf_escape(str(chapter.get("heading") or chapter.get("title") or "未命名章节")), h2_style))
            for paragraph in self._chapter_paragraphs(chapter):
                story.append(Paragraph(self._pdf_escape(self._strip_inline_evidence_ids(paragraph)), base))
            story.extend(self._pdf_evidence_flowables(chapter.get("evidence_ids"), material_map, font_name, small_style, Table, TableStyle, colors))
        story.extend([PageBreak(), Paragraph("参考证据", h2_style)])
        if materials:
            for index, item in enumerate(materials, start=1):
                story.append(
                    self._pdf_info_box(
                        f"来源 {index}：{item.source_title or item.intelligence_title or f'证据 {item.intelligence_id}'}",
                        "\n".join(
                            value
                            for value in [
                                item.quote_text or item.intelligence_summary or "",
                                f"原文：{item.source_url}" if item.source_url else "",
                            ]
                            if value
                        ),
                        font_name,
                        small_style,
                        Table,
                        TableStyle,
                        colors,
                    )
                )
                story.append(Spacer(1, 6))
        else:
            story.append(Paragraph("暂无引用证据。", base))

        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
            title=report.title,
            author="研发营销市场洞察平台",
        )
        doc.build(story)

    def _write_report_docx(self, report: InsightReportDetail, file_path: Path) -> None:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt, RGBColor
        except ImportError as exc:
            raise RuntimeError("缺少 python-docx 依赖，无法生成 DOCX") from exc

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

        font_name = "微软雅黑"
        self._docx_set_style_font(doc.styles["Normal"], font_name, Pt(10.5), RGBColor(52, 72, 70), qn)
        self._docx_set_style_font(doc.styles["Heading 1"], font_name, Pt(16), RGBColor(15, 63, 58), qn)
        self._docx_set_style_font(doc.styles["Heading 2"], font_name, Pt(13), RGBColor(15, 63, 58), qn)

        content = report.content_json or {}
        chapters = content.get("chapters") if isinstance(content, dict) else []
        if not isinstance(chapters, list):
            chapters = []
        materials = report.materials or []
        material_map = {item.intelligence_id: item for item in materials}

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(report.title)
        self._docx_set_run_font(title_run, font_name, Pt(20), RGBColor(16, 32, 31), qn, bold=True)
        self._docx_set_paragraph_spacing(title, after=Pt(10))

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = (
            f"报告类型：{report.report_type}  |  版本：第 {report.version_no} 版  |  "
            f"证据：{report.material_count} 条  |  导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        meta_run = meta.add_run(meta_text)
        self._docx_set_run_font(meta_run, font_name, Pt(9), RGBColor(102, 122, 119), qn)
        self._docx_set_paragraph_spacing(meta, after=Pt(14))

        self._docx_add_info_box(
            doc,
            "摘要",
            self._strip_inline_evidence_ids(str(content.get("executive_summary") or report.summary or "暂无摘要。")),
            font_name,
            qn,
            OxmlElement,
            Pt,
            RGBColor,
        )

        for chapter in chapters:
            if not isinstance(chapter, dict):
                continue
            heading = str(chapter.get("heading") or chapter.get("title") or "未命名章节")
            doc.add_heading(heading, level=1)
            for paragraph_text in self._chapter_paragraphs(chapter):
                self._docx_add_body_paragraph(
                    doc,
                    self._strip_inline_evidence_ids(paragraph_text),
                    font_name,
                    qn,
                    Pt,
                    RGBColor,
                )
            self._docx_add_evidence_boxes(
                doc,
                chapter.get("evidence_ids"),
                material_map,
                font_name,
                qn,
                OxmlElement,
                Pt,
                RGBColor,
            )

        doc.add_page_break()
        doc.add_heading("参考证据", level=1)
        if materials:
            for index, item in enumerate(materials, start=1):
                source_title = item.source_title or item.intelligence_title or f"证据 {item.intelligence_id}"
                body = "\n".join(
                    value
                    for value in [
                        item.quote_text or item.intelligence_summary or "",
                        f"原文：{item.source_url}" if item.source_url else "",
                    ]
                    if value
                )
                self._docx_add_info_box(
                    doc,
                    f"来源 {index}：{source_title}",
                    body,
                    font_name,
                    qn,
                    OxmlElement,
                    Pt,
                    RGBColor,
                    fill="F7FBFA",
                )
        else:
            self._docx_add_body_paragraph(doc, "暂无引用证据。", font_name, qn, Pt, RGBColor)

        doc.core_properties.title = report.title
        doc.core_properties.author = "研发营销市场洞察平台"
        doc.save(str(file_path))

    def _docx_set_style_font(self, style: Any, font_name: str, size: Any, color: Any, qn: Any) -> None:
        style.font.name = font_name
        style.font.size = size
        style.font.color.rgb = color
        style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    def _docx_set_run_font(
        self,
        run: Any,
        font_name: str,
        size: Any,
        color: Any,
        qn: Any,
        *,
        bold: bool = False,
    ) -> None:
        run.font.name = font_name
        run.font.size = size
        run.font.color.rgb = color
        run.bold = bold
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), font_name)

    def _docx_set_paragraph_spacing(self, paragraph: Any, *, before: Any | None = None, after: Any | None = None) -> None:
        if before is not None:
            paragraph.paragraph_format.space_before = before
        if after is not None:
            paragraph.paragraph_format.space_after = after
        paragraph.paragraph_format.line_spacing = 1.35

    def _docx_add_body_paragraph(self, doc: Any, text: str, font_name: str, qn: Any, pt_cls: Any, rgb_cls: Any) -> None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(self._normalize_evidence_text(text))
        self._docx_set_run_font(run, font_name, pt_cls(10.5), rgb_cls(52, 72, 70), qn)
        self._docx_set_paragraph_spacing(paragraph, after=pt_cls(6))

    def _docx_add_info_box(
        self,
        doc: Any,
        title: str,
        body: str,
        font_name: str,
        qn: Any,
        oxml_element_cls: Any,
        pt_cls: Any,
        rgb_cls: Any,
        *,
        fill: str = "F3FBFA",
    ) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.autofit = True
        cell = table.cell(0, 0)
        self._docx_shade_cell(cell, fill, qn, oxml_element_cls)
        title_paragraph = cell.paragraphs[0]
        title_run = title_paragraph.add_run(self._normalize_evidence_text(title))
        self._docx_set_run_font(title_run, font_name, pt_cls(10.5), rgb_cls(15, 118, 110), qn, bold=True)
        self._docx_set_paragraph_spacing(title_paragraph, after=pt_cls(4))
        for line in self._split_pdf_lines(self._normalize_evidence_text(body)):
            paragraph = cell.add_paragraph()
            run = paragraph.add_run(line)
            self._docx_set_run_font(run, font_name, pt_cls(9.5), rgb_cls(80, 100, 97), qn)
            self._docx_set_paragraph_spacing(paragraph, after=pt_cls(3))
        spacer = doc.add_paragraph()
        self._docx_set_paragraph_spacing(spacer, after=pt_cls(6))

    def _docx_shade_cell(self, cell: Any, fill: str, qn: Any, oxml_element_cls: Any) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = oxml_element_cls("w:shd")
        shading.set(qn("w:fill"), fill)
        tc_pr.append(shading)

    def _docx_add_evidence_boxes(
        self,
        doc: Any,
        evidence_ids: Any,
        material_map: dict[int, InsightReportMaterialRead],
        font_name: str,
        qn: Any,
        oxml_element_cls: Any,
        pt_cls: Any,
        rgb_cls: Any,
    ) -> None:
        if not isinstance(evidence_ids, list):
            return
        seen: set[int] = set()
        for raw_id in evidence_ids[:8]:
            try:
                intelligence_id = int(raw_id)
            except (TypeError, ValueError):
                continue
            if intelligence_id in seen:
                continue
            seen.add(intelligence_id)
            item = material_map.get(intelligence_id)
            if not item:
                continue
            self._docx_add_info_box(
                doc,
                f"证据：{item.source_title or item.intelligence_title or f'情报记录 {intelligence_id}'}",
                item.quote_text or item.intelligence_summary or "",
                font_name,
                qn,
                oxml_element_cls,
                pt_cls,
                rgb_cls,
                fill="EEF7FF",
            )

    def _register_pdf_font(self, pdfmetrics: Any, ttfont_cls: Any) -> str:
        font_candidates = [
            Path("C:/Windows/Fonts/NotoSansSC-VF.ttf"),
            Path("C:/Windows/Fonts/msyh.ttc"),
            Path("C:/Windows/Fonts/simhei.ttf"),
            Path("C:/Windows/Fonts/simsun.ttc"),
        ]
        for font_path in font_candidates:
            if font_path.exists():
                font_name = "InsightCJK"
                try:
                    pdfmetrics.getFont(font_name)
                except KeyError:
                    pdfmetrics.registerFont(ttfont_cls(font_name, str(font_path)))
                return font_name
        return "Helvetica"

    def _pdf_info_box(
        self,
        title: str,
        body: str,
        font_name: str,
        paragraph_style: Any,
        table_cls: Any,
        table_style_cls: Any,
        colors_module: Any,
    ) -> Any:
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import Paragraph

        heading_style = ParagraphStyle(
            f"InsightBoxHeading{abs(hash(title))}",
            parent=paragraph_style,
            fontName=font_name,
            fontSize=10,
            leading=15,
            textColor=colors_module.HexColor("#0f766e"),
            spaceAfter=4,
        )
        content = [
            Paragraph(self._pdf_escape(title), heading_style),
            *[Paragraph(self._pdf_escape(part), paragraph_style) for part in self._split_pdf_lines(body)],
        ]
        table = table_cls([[content]], colWidths=[170 * mm])
        table.setStyle(
            table_style_cls(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors_module.HexColor("#f5fbfa")),
                    ("BOX", (0, 0), (-1, -1), 0.5, colors_module.HexColor("#dbe7e5")),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            )
        )
        return table

    def _pdf_evidence_flowables(
        self,
        evidence_ids: Any,
        material_map: dict[int, InsightReportMaterialRead],
        font_name: str,
        paragraph_style: Any,
        table_cls: Any,
        table_style_cls: Any,
        colors_module: Any,
    ) -> list[Any]:
        if not isinstance(evidence_ids, list):
            return []
        flowables: list[Any] = []
        seen: set[int] = set()
        for raw_id in evidence_ids[:8]:
            try:
                intelligence_id = int(raw_id)
            except (TypeError, ValueError):
                continue
            if intelligence_id in seen:
                continue
            seen.add(intelligence_id)
            item = material_map.get(intelligence_id)
            if not item:
                continue
            from reportlab.platypus import Spacer

            flowables.append(
                self._pdf_info_box(
                    f"证据：{item.source_title or item.intelligence_title or f'情报记录 {intelligence_id}'}",
                    item.quote_text or item.intelligence_summary or "",
                    font_name,
                    paragraph_style,
                    table_cls,
                    table_style_cls,
                    colors_module,
                )
            )
            flowables.append(Spacer(1, 5))
        return flowables

    def _chapter_paragraphs(self, chapter: dict[str, Any]) -> list[str]:
        body = chapter.get("body") or chapter.get("content") or chapter.get("paragraphs") or ""
        if isinstance(body, list):
            paragraphs = [self._normalize_evidence_text(str(item).strip()) for item in body if str(item).strip()]
        else:
            paragraphs = [self._normalize_evidence_text(part.strip()) for part in str(body).split("\n") if part.strip()]
        return paragraphs or ["暂无正文。"]

    def _split_pdf_lines(self, text: str) -> list[str]:
        parts = [part.strip() for part in str(text or "").split("\n") if part.strip()]
        return parts or ["暂无内容。"]

    def _pdf_escape(self, text: str) -> str:
        return escape(self._normalize_evidence_text(str(text or ""))).replace("\n", "<br/>")

    def _render_report_html(self, report: InsightReportDetail) -> str:
        content = report.content_json or {}
        chapters = content.get("chapters") if isinstance(content, dict) else []
        if not isinstance(chapters, list):
            chapters = []
        materials = report.materials or []
        material_map = {item.intelligence_id: item for item in materials}
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        chapter_html = "\n".join(self._render_chapter_html(item, material_map) for item in chapters if isinstance(item, dict))
        material_html = "\n".join(
            (
                f'<article class="source-card" id="source-{item.intelligence_id}">'
                f'<div class="source-index">来源 {index}</div>'
                f"<h3>{escape(self._normalize_evidence_text(item.source_title or item.intelligence_title or f'证据 {item.intelligence_id}'))}</h3>"
                f"<p>{escape(self._normalize_evidence_text(item.quote_text or item.intelligence_summary or ''))}</p>"
                f"{self._render_source_link(item.source_url)}"
                "</article>"
            )
            for index, item in enumerate(materials, start=1)
        )
        return f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>{escape(self._normalize_evidence_text(report.title))}</title>
  <style>
    :root {{ color-scheme: light; }}
    * {{ box-sizing: border-box; }}
    body {{ margin: 0; background: #f4f7f8; color: #10201f; font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "Microsoft YaHei", sans-serif; }}
    main {{ width: min(980px, calc(100% - 48px)); margin: 32px auto; background: #fff; border: 1px solid #dbe7e5; border-radius: 16px; padding: 46px; box-shadow: 0 24px 80px rgba(15, 23, 42, .08); }}
    h1 {{ margin: 0; font-size: 30px; line-height: 1.25; letter-spacing: 0; }}
    h2 {{ margin-top: 36px; border-top: 1px solid #dbe7e5; padding-top: 26px; font-size: 21px; letter-spacing: 0; }}
    h3 {{ margin: 0 0 10px; font-size: 16px; letter-spacing: 0; }}
    p {{ line-height: 1.86; color: #344846; font-size: 15px; }}
    a {{ color: #0f766e; word-break: break-word; text-decoration-thickness: 1px; text-underline-offset: 3px; }}
    .meta {{ margin-top: 12px; color: #667a77; font-size: 13px; }}
    .summary {{ margin-top: 28px; padding: 20px 22px; background: #f5fbfa; border-left: 4px solid #0f766e; border-radius: 10px; }}
    .summary strong {{ color: #0f3f3a; }}
    .evidence-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(220px, 1fr)); gap: 12px; margin: 18px 0 8px; }}
    .evidence-card {{ border: 1px solid #dbe7e5; border-radius: 10px; background: #fbfdfc; overflow: hidden; }}
    .evidence-card summary {{ cursor: pointer; list-style: none; padding: 12px 14px; color: #0f766e; font-weight: 650; }}
    .evidence-card summary::-webkit-details-marker {{ display: none; }}
    .evidence-card div {{ border-top: 1px solid #dbe7e5; padding: 12px 14px 14px; }}
    .evidence-card p {{ margin: 8px 0 10px; font-size: 14px; line-height: 1.65; }}
    .source-list {{ display: grid; gap: 14px; margin-top: 18px; }}
    .source-card {{ border: 1px solid #dbe7e5; border-radius: 12px; padding: 16px 18px; background: #fbfdfc; scroll-margin-top: 24px; }}
    .source-card:target {{ border-color: #0f766e; box-shadow: 0 0 0 3px rgba(15, 118, 110, .12); }}
    .source-index {{ margin-bottom: 8px; color: #0f766e; font-size: 12px; font-weight: 700; }}
    @media (max-width: 720px) {{
      main {{ width: calc(100% - 24px); margin: 12px auto; padding: 24px 18px; border-radius: 12px; }}
      h1 {{ font-size: 24px; }}
      h2 {{ font-size: 19px; }}
    }}
  </style>
</head>
<body>
  <main>
    <h1>{escape(self._normalize_evidence_text(report.title))}</h1>
    <div class="meta">报告类型：{escape(report.report_type)} · 版本：第 {report.version_no} 版 · 证据：{report.material_count} 条 · 导出时间：{generated_at}</div>
    <section class="summary">
      <strong>摘要</strong>
      <p>{escape(self._normalize_evidence_text(self._strip_inline_evidence_ids(str(content.get("executive_summary") or report.summary or "暂无摘要。"))))}</p>
    </section>
    {chapter_html}
    <h2>参考证据</h2>
    <div class="source-list">{material_html or "<p>暂无引用证据。</p>"}</div>
  </main>
</body>
</html>
"""

    def _render_chapter_html(self, chapter: dict[str, Any], material_map: dict[int, InsightReportMaterialRead]) -> str:
        heading = escape(self._normalize_evidence_text(str(chapter.get("heading") or chapter.get("title") or "未命名章节")))
        body = chapter.get("body") or chapter.get("content") or chapter.get("paragraphs") or ""
        if isinstance(body, list):
            paragraphs = [str(item) for item in body]
        else:
            paragraphs = [part.strip() for part in str(body).split("\n") if part.strip()]
        paragraph_html = "".join(f"<p>{escape(self._normalize_evidence_text(self._strip_inline_evidence_ids(item)))}</p>" for item in paragraphs) or "<p>暂无正文。</p>"
        evidence_html = self._render_evidence_cards(chapter.get("evidence_ids"), material_map)
        return f"<section><h2>{heading}</h2>{paragraph_html}{evidence_html}</section>"

    def _normalize_evidence_text(self, text: str) -> str:
        return (
            str(text or "")
            .replace("报告素材池", "正式情报资产库")
            .replace("素材池", "资产库")
            .replace("测试素材", "测试数据")
            .replace("证据矩阵", "素材清单")
            .replace("引用证据", "引用素材")
            .replace("报告证据", "报告素材")
        )

    def _strip_inline_evidence_ids(self, text: str) -> str:
        text = re.sub(r"[（(]\s*证据\s*(?:ID|编号|id)?\s*[:：]?\s*[\d,，、\s]+[）)]", "", text)
        text = re.sub(r"证据\s*(?:ID|编号|id)\s*[:：]\s*[\d,，、\s]+", "", text)
        return re.sub(r"\s{2,}", " ", text).strip()

    def _render_evidence_cards(self, evidence_ids: Any, material_map: dict[int, InsightReportMaterialRead]) -> str:
        if not isinstance(evidence_ids, list):
            return ""
        cards: list[str] = []
        seen: set[int] = set()
        for raw_id in evidence_ids:
            try:
                intelligence_id = int(raw_id)
            except (TypeError, ValueError):
                continue
            if intelligence_id in seen:
                continue
            seen.add(intelligence_id)
            item = material_map.get(intelligence_id)
            if not item:
                continue
            title = escape(self._normalize_evidence_text(item.source_title or item.intelligence_title or f"情报记录 {intelligence_id}"))
            summary = escape(self._normalize_evidence_text(item.quote_text or item.intelligence_summary or ""))
            source_link = self._render_source_link(item.source_url)
            cards.append(
                '<details class="evidence-card">'
                f"<summary>来源：{title}</summary>"
                "<div>"
                f"<p>{summary}</p>"
                f'<a href="#source-{intelligence_id}">定位到参考证据</a>'
                f"{' · ' + source_link if source_link else ''}"
                "</div>"
                "</details>"
            )
        if not cards:
            return ""
        return f'<div class="evidence-grid">{"".join(cards)}</div>'

    def _render_source_link(self, source_url: str | None) -> str:
        if not source_url:
            return ""
        safe_url = escape(source_url, quote=True)
        return f'<a href="{safe_url}" target="_blank" rel="noreferrer">查看原文</a>'

    def _to_report_detail(
        self,
        row: InsightReport,
        materials: list[InsightReportMaterialRead],
        versions: list[InsightReportVersionRead],
        charts: list[InsightReportChartRead],
    ) -> InsightReportDetail:
        return InsightReportDetail(**self._to_report_read(row).model_dump(), materials=materials, versions=versions, charts=charts)

    def _to_material_read(self, row: InsightReportMaterial, intelligence: InsightIntelligence | None) -> InsightReportMaterialRead:
        return InsightReportMaterialRead(
            id=row.id or 0,
            create_time=row.create_time,
            update_time=row.update_time,
            report_id=row.report_id,
            intelligence_id=row.intelligence_id,
            section_key=row.section_key,
            sort_no=row.sort_no,
            quote_text=row.quote_text,
            source_url=row.source_url,
            source_title=row.source_title,
            selection_source=row.selection_source,
            selection_reason=row.selection_reason,
            intelligence_title=intelligence.title if intelligence else None,
            intelligence_summary=intelligence.summary if intelligence else None,
        )

    def _to_version_read(self, row: InsightReportVersion) -> InsightReportVersionRead:
        return InsightReportVersionRead(
            id=row.id or 0,
            create_time=row.create_time,
            update_time=row.update_time,
            report_id=row.report_id,
            version_no=row.version_no,
            content_json=row.content_json,
            change_summary=row.change_summary,
            created_by_user_id=row.created_by_user_id,
        )


insight_report_service = InsightReportService()
